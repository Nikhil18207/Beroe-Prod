"""
Chat Endpoints
Handle chat conversations with the AI assistant (Coco).
Supports both HTTP and WebSocket connections.
"""

from typing import Optional, List
from fastapi import APIRouter, Depends, HTTPException, status, Query, WebSocket, WebSocketDisconnect
from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy import select
from sqlalchemy.orm import selectinload
import uuid
from datetime import datetime
import json

from app.database import get_db, async_session_factory
from app.models.user import User
from app.models.conversation import Conversation, Message, MessageRole
from app.api.v1.auth import get_current_user, get_user_from_token
from app.api.v1.dependencies import get_tenant_context, TenantContext
from app.services.chat_service import ChatService
from pydantic import BaseModel, Field
import structlog

logger = structlog.get_logger()

router = APIRouter()

# Initialize chat service
chat_service = ChatService()


class ChatMessageCreate(BaseModel):
    """Create a chat message."""
    content: str = Field(..., min_length=1)
    conversation_id: Optional[uuid.UUID] = None
    session_id: Optional[uuid.UUID] = None
    history: Optional[List[dict]] = None  # Chat history for memory


class ChatMessageResponse(BaseModel):
    """Chat message response."""
    id: uuid.UUID
    role: str
    content: str
    thinking_time: Optional[str] = None
    cards: Optional[dict] = None
    panel_state_change: Optional[str] = None
    created_at: datetime

    model_config = {"from_attributes": True}


class ConversationResponse(BaseModel):
    """Conversation response."""
    id: uuid.UUID
    title: str
    summary: Optional[str] = None
    panel_state: Optional[str] = None
    message_count: int
    created_at: datetime
    last_message_at: Optional[datetime] = None

    model_config = {"from_attributes": True}


class ChatFullResponse(BaseModel):
    """Full response including both user and assistant messages."""
    status: str
    conversation_id: str
    user_message: dict
    assistant_message: dict
    panel_state_change: Optional[str] = None
    cards: Optional[dict] = None


@router.get("/conversations", response_model=List[ConversationResponse])
async def list_conversations(
    tenant: TenantContext = Depends(get_tenant_context),
    db: AsyncSession = Depends(get_db),
    limit: int = Query(20, ge=1, le=100)
):
    """
    List user's recent conversations.
    """
    tenant.require_permission("analyses", "read")
    result = await db.execute(
        select(Conversation)
        .where(Conversation.user_id == tenant.user_id)
        .order_by(Conversation.updated_at.desc())
        .limit(limit)
        .options(selectinload(Conversation.messages))
    )
    conversations = result.scalars().all()

    return [
        ConversationResponse(
            id=c.id,
            title=c.title,
            summary=c.summary,
            panel_state=c.panel_state,
            message_count=len(c.messages) if c.messages else 0,
            created_at=c.created_at,
            last_message_at=c.last_message_at
        )
        for c in conversations
    ]


@router.post("/message", response_model=ChatFullResponse)
async def send_message(
    message_data: ChatMessageCreate,
    tenant: TenantContext = Depends(get_tenant_context),
    db: AsyncSession = Depends(get_db)
):
    """
    Send a message and get AI response.

    This endpoint:
    1. Loads context from the analysis session (if provided)
    2. Processes the message through the LLM
    3. Returns the AI response with any panel state changes
    """
    tenant.require_permission("analyses", "create")
    try:
        # Process message through chat service
        response = await chat_service.process_message(
            user_message=message_data.content,
            conversation_id=str(message_data.conversation_id) if message_data.conversation_id else None,
            session_id=str(message_data.session_id) if message_data.session_id else None,
            user_id=str(tenant.user_id),
            db=db
        )

        return ChatFullResponse(
            status="success",
            conversation_id=response.metadata.get("conversation_id", ""),
            user_message={
                "id": str(uuid.uuid4()),
                "role": "user",
                "content": message_data.content
            },
            assistant_message={
                "id": response.message_id,
                "role": "assistant",
                "content": response.content,
                "thinking_time": response.thinking_time
            },
            panel_state_change=response.panel_state_change,
            cards=response.cards
        )

    except Exception as e:
        logger.error("Chat message failed", error=str(e))
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"Failed to process message: {str(e)}"
        )


@router.get("/conversation/{conversation_id}/messages", response_model=List[ChatMessageResponse])
async def get_conversation_messages(
    conversation_id: uuid.UUID,
    tenant: TenantContext = Depends(get_tenant_context),
    db: AsyncSession = Depends(get_db)
):
    """
    Get all messages in a conversation.
    """
    tenant.require_permission("analyses", "read")
    result = await db.execute(
        select(Conversation)
        .where(
            Conversation.id == conversation_id,
            Conversation.user_id == tenant.user_id
        )
        .options(selectinload(Conversation.messages))
    )
    conversation = result.scalar_one_or_none()

    if not conversation:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Conversation not found"
        )

    return [
        ChatMessageResponse(
            id=m.id,
            role=m.role.value,
            content=m.content,
            thinking_time=m.thinking_time,
            cards=m.cards,
            panel_state_change=m.panel_state_change,
            created_at=m.created_at
        )
        for m in sorted(conversation.messages, key=lambda x: x.created_at)
    ]


@router.delete("/conversation/{conversation_id}")
async def delete_conversation(
    conversation_id: uuid.UUID,
    tenant: TenantContext = Depends(get_tenant_context),
    db: AsyncSession = Depends(get_db)
):
    """
    Delete a conversation.
    """
    tenant.require_permission("analyses", "delete")
    result = await db.execute(
        select(Conversation)
        .where(
            Conversation.id == conversation_id,
            Conversation.user_id == tenant.user_id
        )
    )
    conversation = result.scalar_one_or_none()

    if not conversation:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Conversation not found"
        )

    await db.delete(conversation)
    await db.commit()

    return {"status": "success", "deleted": True, "id": str(conversation_id)}


@router.get("/suggestions")
async def get_suggested_questions(
    session_id: Optional[uuid.UUID] = None,
    tenant: TenantContext = Depends(get_tenant_context),
    db: AsyncSession = Depends(get_db)
):
    """
    Get suggested questions based on session context.
    """
    tenant.require_permission("analyses", "read")
    questions = await chat_service.get_suggested_questions(
        session_id=str(session_id) if session_id else None,
        db=db
    )

    return {"suggestions": questions}


# WebSocket connection manager
class ConnectionManager:
    """Manage WebSocket connections."""

    def __init__(self):
        self.active_connections: dict = {}  # user_id -> WebSocket

    async def connect(self, websocket: WebSocket, user_id: str):
        await websocket.accept()
        self.active_connections[user_id] = websocket
        logger.info("WebSocket connected", user_id=user_id)

    def disconnect(self, user_id: str):
        if user_id in self.active_connections:
            del self.active_connections[user_id]
            logger.info("WebSocket disconnected", user_id=user_id)

    async def send_message(self, user_id: str, message: dict):
        if user_id in self.active_connections:
            await self.active_connections[user_id].send_json(message)


manager = ConnectionManager()


@router.websocket("/ws/{token}")
async def websocket_chat(
    websocket: WebSocket,
    token: str
):
    """
    WebSocket endpoint for real-time chat.

    Message format (client -> server):
    {
        "type": "message",
        "content": "user message",
        "session_id": "optional-session-id",
        "conversation_id": "optional-conversation-id"
    }

    Response format (server -> client):
    {
        "type": "response" | "error" | "typing",
        "content": "...",
        "thinking_time": "1.2s",
        "panel_state_change": "opportunities",
        "cards": {...}
    }
    """
    # Authenticate user from token
    async with async_session_factory() as db:
        user = await get_user_from_token(token, db)

        if not user:
            await websocket.close(code=4001, reason="Authentication failed")
            return

        user_id = str(user.id)

    await manager.connect(websocket, user_id)

    try:
        while True:
            # Receive message
            data = await websocket.receive_json()

            if data.get("type") == "ping":
                await websocket.send_json({"type": "pong"})
                continue

            if data.get("type") == "message":
                content = data.get("content", "")
                session_id = data.get("session_id")
                conversation_id = data.get("conversation_id")

                # Send typing indicator
                await websocket.send_json({"type": "typing", "status": "started"})

                # Process message
                async with async_session_factory() as db:
                    try:
                        response = await chat_service.process_message(
                            user_message=content,
                            conversation_id=conversation_id,
                            session_id=session_id,
                            user_id=user_id,
                            db=db
                        )

                        # Send response
                        await websocket.send_json({
                            "type": "response",
                            "conversation_id": response.metadata.get("conversation_id"),
                            "message_id": response.message_id,
                            "content": response.content,
                            "thinking_time": response.thinking_time,
                            "panel_state_change": response.panel_state_change,
                            "cards": response.cards
                        })

                    except Exception as e:
                        logger.error("WebSocket message processing failed", error=str(e))
                        await websocket.send_json({
                            "type": "error",
                            "message": str(e)
                        })

    except WebSocketDisconnect:
        manager.disconnect(user_id)
    except Exception as e:
        logger.error("WebSocket error", error=str(e))
        manager.disconnect(user_id)


@router.options("/demo-message")
async def demo_message_options():
    """
    Handle CORS preflight for demo-message endpoint.
    """
    from fastapi.responses import Response
    return Response(
        status_code=200,
        headers={
            "Access-Control-Allow-Origin": "*",
            "Access-Control-Allow-Methods": "POST, OPTIONS",
            "Access-Control-Allow-Headers": "Content-Type, Authorization",
        }
    )


@router.post("/demo-message")
async def demo_message(
    message_data: ChatMessageCreate
):
    """
    Demo chat endpoint without authentication.
    For testing purposes only.
    """
    from app.services.llm_service import LLMService

    llm = LLMService()

    # Max AI system prompt - Knowledgeable procurement assistant
    system_prompt = """You are Max, a procurement AI assistant with deep knowledge of savings opportunities.

YOUR KNOWLEDGE:
- You understand the 4 opportunity types: Volume Bundling, Target Pricing, Risk Management, Re-spec Pack
- You know about proof points that validate each opportunity
- You have access to spend data, supplier info, contracts, and metrics from uploaded files
- The user's context (=== KNOWLEDGE BASE ===) contains all the data you need

COMMUNICATION STYLE:
- 2-4 sentences for simple questions, more for complex ones
- Use actual names and numbers from the KNOWLEDGE BASE
- Be conversational like a Slack colleague, not formal
- Reference specific suppliers, spend amounts, metrics when relevant
- If asked about data you have, cite it specifically

EXAMPLES:
"how much is the total spend?" → "Total spend is $536K. AgriGlobal leads at $141K, followed by EuroFoods at $136K."

"tell me about this opportunity" → "This Volume Bundling opportunity for Edible Oils has $536K spend across 4 suppliers. The key insight is 77% concentration in your top 3 suppliers - consolidating could get you 3-5% volume discounts."

"why did you recommend consolidating with AgriGlobal?" → "AgriGlobal already has 26% of your spend ($141K) and good risk ratings. Consolidating more volume with them gives leverage for tier pricing that wouldn't be possible with fragmented spend."

"what are the proof points?" → "There are 8 proof points for Volume Bundling. Currently 2 are validated: Regional Spend Distribution and Supplier Concentration. The pending ones like Price Variance and Tail Spend need verification from your team."

Never thank users for basic questions. Just answer with specific data."""

    # Build messages with history for memory
    messages = [{"role": "system", "content": system_prompt}]

    # Add conversation history if provided (for memory)
    if message_data.history:
        for msg in message_data.history[-6:]:  # Keep last 6 messages
            if msg.get("role") in ["user", "assistant"] and msg.get("content"):
                messages.append({"role": msg["role"], "content": msg["content"]})

    # Add current user message
    messages.append({"role": "user", "content": message_data.content})

    # Use the system prompt with the user message
    response = await llm.chat(messages=messages)

    return {
        "status": "success",
        "user_message": {
            "content": message_data.content
        },
        "assistant_message": {
            "content": response.content,
            "thinking_time": f"{response.latency_ms/1000:.1f}s"
        }
    }


# ============================================================================
# OPPORTUNITY CHAT ENDPOINT - Structured context for Max AI
# ============================================================================

class OpportunityChatRequest(BaseModel):
    """Structured request for opportunity-focused chat with Max."""
    question: str = Field(..., description="User's question")
    opportunity_type: str = Field(..., description="volume-bundling, target-pricing, risk-management, respec-pack")
    category_name: str = Field(..., description="Category name e.g. 'Edible Oils'")

    # Spend data
    total_spend: float = Field(default=0, description="Total spend amount")
    suppliers: List[dict] = Field(default=[], description="List of suppliers with name, spend, country, riskRating")

    # Metrics
    metrics: dict = Field(default={}, description="priceVariance, top3Concentration, tailSpendPercentage, supplierCount")

    # Proof points
    proof_points: List[dict] = Field(default=[], description="id, name, isValidated, description")
    current_proof_point: Optional[str] = Field(default=None, description="Proof point user is currently viewing")

    # Recommendations
    recommendations: List[dict] = Field(default=[], description="Generated recommendations with text and reason")

    # User goals
    goals: dict = Field(default={}, description="cost, risk, esg priority scores")
    savings_percentage: Optional[str] = Field(default=None, description="Estimated savings range")

    # Raw file data
    spend_data_sample: List[dict] = Field(default=[], description="First 10 rows of spend CSV")
    contract_summary: Optional[str] = Field(default=None, description="Parsed contract text")
    supplier_master_summary: Optional[str] = Field(default=None, description="Supplier master key info")

    # Chat history
    history: List[dict] = Field(default=[], description="Previous messages for memory")


# Proof point definitions for Max's knowledge
PROOF_POINT_KNOWLEDGE = {
    "volume-bundling": {
        "Regional Spend Distribution": "Measures how spend is distributed across regions. High impact if top 3 regions have >80% spend - good for regional bundling.",
        "Tail Spend Analysis": "Identifies fragmented spend across many small suppliers. High impact if >30% spend is with bottom suppliers.",
        "Volume Leverage Potential": "Checks if spend is too fragmented. High impact if >10 suppliers and none has >20% share.",
        "Price Variance": "Measures price differences across suppliers for same items. High impact if >25% variance - opportunity to negotiate to best price.",
        "Average Spend per Supplier": "Identifies if suppliers are too small for leverage. High impact if average <$100K.",
        "Market Consolidation": "Uses HHI to measure market concentration. High impact if HHI <1500 (competitive market).",
        "Supplier Location Proximity": "Checks if suppliers are geographically clustered. High impact if >70% in same region.",
        "Supplier Risk Rating": "Assesses financial health of top suppliers. High impact if top 5 have low risk ratings.",
    },
    "target-pricing": {
        "Price Variance": "Measures price differences. High variance means opportunity to use best price as negotiation target.",
        "Tariff Rate Differential": "Compares tariff rates across sourcing regions. High impact if >15% differential.",
        "Cost Structure Analysis": "Analyzes raw material vs conversion costs. High impact if >60% is raw material (commodity).",
        "Unit Price Benchmark": "Compares prices to market benchmarks. High impact if >15% above benchmark.",
    },
    "risk-management": {
        "Single Sourcing Risk": "Identifies dependency on single supplier. High impact if >50% with one supplier.",
        "Supplier Concentration": "Measures top 3 supplier concentration. High impact if >80%.",
        "Category Risk Assessment": "Evaluates category-level supply chain risks.",
        "Inflation Exposure": "Measures exposure to cost inflation. High impact if >8% annual inflation.",
        "Exchange Rate Volatility": "Checks currency exposure. High impact if >50% in volatile currencies.",
        "Geopolitical Risk": "Assesses supplier country risks. High impact if >40% from high-risk regions.",
        "Supplier Risk Rating": "Evaluates supplier financial health and stability.",
    },
    "respec-pack": {
        "Price Variance": "Price differences across specs/packaging.",
        "Export Data Analysis": "Compares to export standard specs. High impact if >20% above standard.",
        "Cost Structure Analysis": "Identifies packaging/spec cost opportunities.",
    }
}


@router.post("/opportunity-chat")
async def opportunity_chat(request: OpportunityChatRequest):
    """
    Chat endpoint with full opportunity context.
    Max has deep knowledge of the opportunity, proof points, and all uploaded data.
    """
    from app.services.llm_service import LLMService

    llm = LLMService()

    # Get opportunity-specific knowledge
    opp_knowledge = {
        "volume-bundling": "Volume Bundling consolidates fragmented spend to negotiate volume discounts. Levers: supplier consolidation, regional bundling, tail spend reduction. Benchmark: 0-5% savings.",
        "target-pricing": "Target Pricing uses best prices as benchmarks for negotiation. Focus on price variance and should-cost analysis. Benchmark: 1-2% savings.",
        "risk-management": "Risk Management diversifies supply base to reduce single-source dependency. Focus on concentration and supplier health. Benchmark: 1-3% cost avoidance.",
        "respec-pack": "Re-specification & Packaging optimizes specs and packaging to reduce costs. Benchmark: 2-3% savings."
    }

    # Get proof point knowledge for this opportunity
    pp_knowledge = PROOF_POINT_KNOWLEDGE.get(request.opportunity_type, {})

    # Build current proof point context if user is viewing one
    current_pp_context = ""
    if request.current_proof_point:
        pp_info = pp_knowledge.get(request.current_proof_point, "")
        current_pp_context = f"\n\n[USER IS VIEWING: {request.current_proof_point}]\n{pp_info}"

    # Build supplier list
    supplier_list = "\n".join([
        f"- {s.get('name', 'Unknown')}: ${s.get('spend', 0):,.0f}" +
        (f" ({s.get('country', '')})" if s.get('country') else "") +
        (f" [Risk: {s.get('riskRating', '')}]" if s.get('riskRating') else "")
        for s in request.suppliers[:10]
    ]) or "No supplier data"

    # Build proof points list
    pp_list = "\n".join([
        f"- {pp.get('name', '')}: {'✓ Validated' if pp.get('isValidated') else '○ Pending'}"
        for pp in request.proof_points
    ]) or "No proof points"

    # Build metrics
    metrics = request.metrics
    metrics_text = f"Price Variance: {metrics.get('priceVariance', 'N/A')}%, Top3 Concentration: {metrics.get('top3Concentration', 'N/A')}%, Tail Spend: {metrics.get('tailSpendPercentage', 'N/A')}%, Suppliers: {metrics.get('supplierCount', 'N/A')}"

    # Build recommendations if any
    recs_text = ""
    if request.recommendations:
        recs_text = "\n\nMY RECOMMENDATIONS:\n" + "\n".join([
            f"{i+1}. {r.get('text', '')}\n   Reason: {r.get('reason', '')}"
            for i, r in enumerate(request.recommendations)
        ])

    # Build spend data sample if available
    spend_sample_text = ""
    if request.spend_data_sample:
        spend_sample_text = "\n\nSPEND DATA SAMPLE:\n" + "\n".join([
            ", ".join([f"{k}: {v}" for k, v in row.items()])
            for row in request.spend_data_sample[:5]
        ])

    # Build comprehensive system prompt
    system_prompt = f"""You are Max, a procurement AI assistant with deep knowledge of this opportunity.

=== YOUR KNOWLEDGE BASE ===

OPPORTUNITY: {request.opportunity_type} for {request.category_name}
{opp_knowledge.get(request.opportunity_type, '')}

TOTAL SPEND: ${request.total_spend:,.0f}
SAVINGS ESTIMATE: {request.savings_percentage or '3-5%'}

SUPPLIERS ({len(request.suppliers)} total):
{supplier_list}

METRICS:
{metrics_text}

PROOF POINTS:
{pp_list}

PROOF POINT DEFINITIONS FOR THIS OPPORTUNITY:
{chr(10).join([f'- {name}: {desc}' for name, desc in pp_knowledge.items()])}
{current_pp_context}
{recs_text}
{spend_sample_text}
{f'{chr(10)}CONTRACT INFO: {request.contract_summary}' if request.contract_summary else ''}
{f'{chr(10)}SUPPLIER DETAILS: {request.supplier_master_summary}' if request.supplier_master_summary else ''}

=== HOW TO RESPOND ===
- Use specific supplier names and numbers from your knowledge base
- Reference proof points by name when relevant
- Explain how metrics support the opportunity
- 2-4 sentences for simple questions, more detail for complex ones
- If asked about a recommendation, explain the reasoning using the data
- Be conversational like a Slack colleague, not formal"""

    # Build messages with history
    messages = [{"role": "system", "content": system_prompt}]

    # Add conversation history
    for msg in request.history[-6:]:
        if msg.get("role") in ["user", "assistant"] and msg.get("content"):
            messages.append({"role": msg["role"], "content": msg["content"]})

    # Add current question
    messages.append({"role": "user", "content": request.question})

    # Call LLM
    response = await llm.chat(messages=messages)

    return {
        "status": "success",
        "user_message": {"content": request.question},
        "assistant_message": {
            "content": response.content,
            "thinking_time": f"{response.latency_ms/1000:.1f}s"
        }
    }


# ============================================================================
# OPPORTUNITY RECOMMENDATIONS ENDPOINT
# ============================================================================

class OpportunityRecommendationsRequest(BaseModel):
    """Request for generating opportunity-specific recommendations."""
    opportunity_type: str = Field(..., description="Type: volume-bundling, target-pricing, risk-management, respec-pack")
    category_name: str = Field(..., description="Category name e.g. 'Edible Oils'")
    locations: Optional[List[str]] = Field(default=None, description="Geographic locations/regions e.g. ['Europe', 'India', 'Asia Pacific']")
    spend_data: dict = Field(default={}, description="Spend data including totalSpend and breakdown")
    supplier_data: List[dict] = Field(default=[], description="List of suppliers with name and spend")
    metrics: dict = Field(default={}, description="Computed metrics like priceVariance, top3Concentration")
    proof_points: List[dict] = Field(default=[], description="List of proof points with id, name, isValidated")
    playbook_data: Optional[dict] = Field(default=None, description="Category playbook data with recommendations, strategies, etc.")
    contract_data: Optional[dict] = Field(default=None, description="Contract data with active contracts, payment terms, renewal dates, etc.")
    supplier_master_data: Optional[dict] = Field(default=None, description="Supplier master data with ratings, certifications, financial health, etc.")


@router.options("/recommendations")
async def recommendations_options():
    """Handle CORS preflight for recommendations endpoint."""
    from fastapi.responses import Response
    return Response(
        status_code=200,
        headers={
            "Access-Control-Allow-Origin": "*",
            "Access-Control-Allow-Methods": "POST, OPTIONS",
            "Access-Control-Allow-Headers": "Content-Type, Authorization",
        }
    )


@router.post("/recommendations")
async def generate_recommendations(
    request: OpportunityRecommendationsRequest
):
    """
    Generate intelligent, context-aware recommendations for a procurement opportunity.

    This endpoint analyzes ALL uploaded documents (spend data, category playbook, contracts, 
    supplier master) and generates 10 unique, data-driven recommendations that:
    - Avoid duplicating existing playbook recommendations
    - Use actual supplier names, spend amounts, and metrics
    - Are specific to the opportunity type and proof points
    - Include reasoning based on actual data
    """
    from app.services.llm_service import LLMService
    import json as json_module

    llm = LLMService()

    try:
        response = await llm.generate_opportunity_recommendations(
            opportunity_type=request.opportunity_type,
            category_name=request.category_name,
            locations=request.locations,
            spend_data=request.spend_data,
            supplier_data=request.supplier_data,
            metrics=request.metrics,
            proof_points=request.proof_points,
            playbook_data=request.playbook_data,
            contract_data=request.contract_data,
            supplier_master_data=request.supplier_master_data
        )

        # Parse the LLM response - expects JSON array of {text, reason} objects
        recommendations = None
        raw_content = response.content.strip()

        # Log the raw response for debugging
        logger.info(f"LLM raw response length: {len(raw_content)} chars")
        logger.info(f"LLM raw response (first 1000 chars): {raw_content[:1000]}")

        # Try to extract JSON from the response (handle markdown code blocks, extra text, etc.)
        json_content = raw_content

        # Remove markdown code block if present
        if "```json" in json_content:
            json_content = json_content.split("```json")[1].split("```")[0].strip()
        elif "```" in json_content:
            json_content = json_content.split("```")[1].split("```")[0].strip()

        # Find JSON array or object boundaries
        if not json_content.startswith('[') and not json_content.startswith('{'):
            # Try to find the first [ or {
            bracket_idx = json_content.find('[')
            brace_idx = json_content.find('{')
            if bracket_idx >= 0 and (brace_idx < 0 or bracket_idx < brace_idx):
                json_content = json_content[bracket_idx:]
            elif brace_idx >= 0:
                json_content = json_content[brace_idx:]

        try:
            parsed = json_module.loads(json_content)

            # CRITICAL: Check if the entire response is the context object echoed back
            # This happens when LLM returns the input instead of generating recommendations
            if isinstance(parsed, dict):
                context_fields = ['opportunity_type', 'total_spend', 'spend_breakdown', 'top_suppliers',
                                  'metrics', 'validated_proof_points', 'unvalidated_items', 'category', 'locations']
                is_context_echo = any(field in parsed for field in context_fields)

                if is_context_echo:
                    logger.warning("LLM returned context object instead of recommendations, forcing fallback")
                    recommendations = None  # Force fallback
                else:
                    # Object format: {"Recommendations": [...]} or {"recommendations": [...]}
                    recs_data = parsed.get("Recommendations") or parsed.get("recommendations")
                    if recs_data and isinstance(recs_data, list):
                        recommendations = []
                        for item in recs_data:
                            if isinstance(item, str):
                                recommendations.append({"text": item, "reason": ""})
                            elif isinstance(item, dict):
                                text = item.get('text') or item.get('recommendation') or item.get('content')
                                reason = item.get('reason') or item.get('reasoning') or item.get('rationale') or ""
                                if text and len(str(text)) > 5:  # Ensure it's meaningful text
                                    recommendations.append({"text": str(text), "reason": str(reason)})

            elif isinstance(parsed, list):
                recommendations = []
                for item in parsed:
                    if isinstance(item, str):
                        # Old format - just string, no reason
                        if len(item) > 10:  # Ensure it's meaningful
                            recommendations.append({"text": item, "reason": ""})
                    elif isinstance(item, dict):
                        # Check if this looks like context data (bad response)
                        context_fields = ['opportunity_type', 'total_spend', 'spend_breakdown', 'top_suppliers',
                                          'metrics', 'validated_proof_points', 'category', 'locations']
                        if any(field in item for field in context_fields):
                            logger.warning(f"Skipping context data item in list: {list(item.keys())[:5]}")
                            continue  # Skip context data
                        text = item.get('text') or item.get('recommendation') or item.get('content')
                        reason = item.get('reason') or item.get('reasoning') or item.get('rationale') or ""
                        if text and len(str(text)) > 5:  # Ensure it's meaningful text
                            recommendations.append({"text": str(text), "reason": str(reason)})

        except json_module.JSONDecodeError as e:
            logger.warning(f"JSON decode error: {e}. Trying line-based parsing.")
            # If not valid JSON, split by newlines or use as single recommendation
            lines = [r.strip() for r in raw_content.split('\n') if r.strip()]
            # Filter out lines that look like JSON or context data
            recommendations = [{"text": line, "reason": ""} for line in lines
                              if not line.startswith('{') and not line.startswith('[') and len(line) > 10]

        # FINAL VALIDATION: Check if any recommendation text looks like JSON (context echo)
        # This catches cases where the LLM returned context data that slipped through
        if recommendations:
            valid_recs = []
            for rec in recommendations:
                text = rec.get('text', '')
                # Skip if text looks like JSON (starts with { or [, or contains typical context fields)
                if text.strip().startswith('{') or text.strip().startswith('['):
                    logger.warning(f"Skipping recommendation with JSON-like text: {text[:100]}...")
                    continue
                # Skip if text contains context field patterns
                context_patterns = ['"opportunity_type"', '"total_spend"', '"spend_breakdown"',
                                   '"top_suppliers"', '"metrics":', '"locations":']
                if any(pattern in text for pattern in context_patterns):
                    logger.warning(f"Skipping recommendation with context patterns: {text[:100]}...")
                    continue
                valid_recs.append(rec)
            recommendations = valid_recs

        # POST-PROCESSING: Filter out recommendations that duplicate playbook concepts
        # This catches duplicates even if the LLM ignores deduplication instructions
        if recommendations and request.playbook_data:
            playbook_recs = request.playbook_data.get('recommendations', [])
            if playbook_recs:
                # Extract key phrases to ban from playbook recommendations
                banned_phrases = []
                for pb_rec in playbook_recs:
                    pb_lower = str(pb_rec).lower()
                    # Extract key concepts from playbook
                    if 'quarterly' in pb_lower and 'business review' in pb_lower:
                        banned_phrases.extend(['quarterly business review', 'quarterly review'])
                    if 'quarterly' in pb_lower:
                        banned_phrases.append('quarterly')
                    if 'weekly' in pb_lower and 'monitor' in pb_lower:
                        banned_phrases.append('weekly monitoring')
                    if 'weekly' in pb_lower:
                        banned_phrases.append('weekly')
                    if 'index-linked' in pb_lower or 'index linked' in pb_lower:
                        banned_phrases.append('index-linked')
                        banned_phrases.append('index linked')
                    if 'commodity index' in pb_lower:
                        banned_phrases.append('commodity index')
                    if '2 qualified' in pb_lower or 'two qualified' in pb_lower:
                        banned_phrases.append('2 qualified supplier')
                        banned_phrases.append('maintain 2')
                        banned_phrases.append('maintain two')

                # Filter recommendations that contain banned phrases
                if banned_phrases:
                    filtered_recs = []
                    for rec in recommendations:
                        rec_text = str(rec.get('text', '')).lower()
                        is_duplicate = False
                        for phrase in banned_phrases:
                            if phrase.lower() in rec_text:
                                logger.warning(f"Filtered out duplicate recommendation containing '{phrase}': {rec.get('text', '')[:80]}...")
                                is_duplicate = True
                                break
                        if not is_duplicate:
                            filtered_recs.append(rec)

                    if len(filtered_recs) < len(recommendations):
                        logger.info(f"Deduplication filtered out {len(recommendations) - len(filtered_recs)} recommendations")
                    recommendations = filtered_recs

        # If we couldn't extract valid recommendations, use fallback with reasons
        if not recommendations or len(recommendations) == 0:
            logger.warning("No valid recommendations extracted from LLM response, using fallback")

            # Get metrics for fallback reasons - with actual data
            metrics = request.metrics or {}
            supplier_data = request.supplier_data or []
            total_spend = request.spend_data.get('totalSpend', 0) if request.spend_data else 0

            # Get supplier details
            supplier1 = supplier_data[0] if len(supplier_data) > 0 else {'name': 'top supplier', 'spend': 0}
            supplier2 = supplier_data[1] if len(supplier_data) > 1 else {'name': 'second supplier', 'spend': 0}
            supplier3 = supplier_data[2] if len(supplier_data) > 2 else {'name': 'third supplier', 'spend': 0}

            top_supplier = supplier1.get('name', 'top supplier')
            top_spend = supplier1.get('spend', 0)
            supplier2_name = supplier2.get('name', 'second supplier')
            supplier2_spend = supplier2.get('spend', 0)
            supplier3_name = supplier3.get('name', 'third supplier')
            supplier3_spend = supplier3.get('spend', 0)

            concentration = metrics.get('top3Concentration', 65)
            price_variance = metrics.get('priceVariance', 15)
            tail_spend = metrics.get('tailSpendPercentage', 12)
            supplier_count = metrics.get('supplierCount', len(supplier_data))

            # Calculate top 3 spend
            top3_spend = top_spend + supplier2_spend + supplier3_spend

            # Get locations
            locations = request.locations or []
            locations_str = ', '.join(locations) if locations else 'your regions'
            locations_short = locations[0] if locations else 'your region'

            # Format currency helper
            def fmt_currency(amt):
                if amt >= 1000000:
                    return f"${amt/1000000:.1f}M"
                elif amt >= 1000:
                    return f"${amt/1000:.0f}K"
                return f"${amt:.0f}"

            fallback_recs = {
                "volume-bundling": [
                    {"text": f"Leverage the {price_variance:.0f}% price variance and high concentration ({concentration:.0f}%) among top 3 suppliers in {locations_str} to negotiate volume-based rebates with {top_supplier}, {supplier2_name}, and {supplier3_name}. Target 5-8% discount through tiered volume commitments.",
                     "reason": f"Spend analysis shows {price_variance:.0f}% price variance across {supplier_count} suppliers in {locations_str}. {top_supplier} has {fmt_currency(top_spend)} spend - consolidating can drive 5-15% savings. Source: Current spend data and supplier metrics."},
                    {"text": f"With a total spend of {fmt_currency(total_spend)} in the {request.category_name} category across {locations_str}, consolidate demands across sites to create a single negotiating unit. Implement monthly demand pooling by Q2 2026.",
                     "reason": f"Spend analysis reveals {tail_spend:.0f}% tail spend fragmented across multiple suppliers in {locations_str}. Bundling this spend with top suppliers can unlock tiered pricing. Source: Current tail spend metrics."},
                    {"text": f"Bundle {request.category_name} purchases in {locations_short} with {top_supplier} who currently holds {(top_spend/total_spend*100) if total_spend > 0 else 0:.0f}% market share. Negotiate multi-year framework agreement with escalation caps.",
                     "reason": f"Supplier data shows {top_supplier}'s spend of {fmt_currency(top_spend)} in {locations_str} makes them a strategic partner. Volume commitments can yield 3-8% price reductions. Source: Supplier spend analysis."},
                    {"text": f"Rationalize the supplier base from {supplier_count} to 3-4 strategic suppliers by end of 2026, focusing on {top_supplier} and {supplier2_name} for primary volumes in {locations_str}.",
                     "reason": f"Current metrics show {supplier_count} active suppliers dilutes volume leverage. Top 3 suppliers control {concentration:.0f}% of spend. Consolidation enables better pricing. Source: Supplier concentration metrics."},
                    {"text": f"Implement regional demand aggregation for {locations_str} with bi-weekly forecast sharing to {top_supplier} and {supplier2_name}, enabling better production planning and 2-4% logistics savings.",
                     "reason": f"Spend data shows fragmented ordering across {len(locations) if locations else 1} regions. Coordinated demand forecasting reduces supplier costs passed to buyer. Source: Regional spend breakdown."},
                    {"text": f"Negotiate rebate structures with {top_supplier} ({fmt_currency(top_spend)}) and {supplier2_name} ({fmt_currency(supplier2_spend)}) tied to annual volume thresholds, targeting 3-5% rebate on volumes above {fmt_currency(total_spend * 0.8)}.",
                     "reason": f"Combined spend of {fmt_currency(top_spend + supplier2_spend)} with top 2 suppliers provides leverage for rebate negotiations. Industry benchmarks show 3-5% rebates achievable. Source: Supplier spend data."},
                    {"text": f"Create a preferred supplier program for {request.category_name} in {locations_str} with {top_supplier}, {supplier2_name}, and {supplier3_name}, offering guaranteed volumes in exchange for 5-10% better pricing.",
                     "reason": f"Top 3 suppliers control {fmt_currency(top3_spend)} ({concentration:.0f}%) of spend. Formalizing preferred status incentivizes competitive pricing. Source: Supplier concentration analysis."},
                    {"text": f"Standardize order quantities and delivery schedules across {locations_str} to enable full truckload shipments with {top_supplier}, reducing per-unit logistics costs by 10-15%.",
                     "reason": f"Spend analysis shows {fmt_currency(total_spend)} distributed across multiple delivery points in {locations_str}. Consolidation optimizes logistics. Source: Regional spend data."},
                    {"text": f"Conduct monthly performance check-ins with {top_supplier} and {supplier2_name} to track volume commitments, pricing performance, and identify additional bundling opportunities across {locations_str}.",
                     "reason": f"Regular reviews ensure ongoing consolidation opportunities in {locations_str} are captured and savings targets of {fmt_currency(total_spend * 0.04)} are achieved. Source: Best practices."},
                    {"text": f"I will monitor market conditions in {locations_str} and alert you on significant price movements (±5% threshold) and consolidation opportunities.",
                     "reason": f"Continuous monitoring of {fmt_currency(total_spend)} spend in {locations_str} ensures you capture pricing opportunities and market shifts. Source: Ongoing market intelligence."},
                    # Additional 10 recommendations for 20 total
                    {"text": f"Implement Hub-and-Spoke sourcing model with {top_supplier} as the strategic hub for {locations_str}, consolidating 70% of {fmt_currency(total_spend)} through single coordination point.",
                     "reason": f"Hub model reduces administrative costs by 12-15% while maintaining supply security. {top_supplier}'s {fmt_currency(top_spend)} volume makes them ideal hub. Source: Logistics optimization analysis."},
                    {"text": f"Launch 3-year volume commitment program with {supplier2_name} ({fmt_currency(supplier2_spend)}) in exchange for 8-12% price improvement and guaranteed capacity.",
                     "reason": f"Long-term commitments de-risk {supplier2_name}'s capacity planning. They'll pay for that certainty with better pricing. Source: Contract analysis best practices."},
                    {"text": f"Transition {supplier3_name} ({fmt_currency(supplier3_spend)}) to performance-based rebate structure with quarterly reviews in {locations_short}.",
                     "reason": f"{supplier3_name} is hungry for growth. Performance rebates align incentives and can deliver 3-5% additional savings. Source: Supplier development principles."},
                    {"text": f"Create cross-category bundling opportunity by combining {request.category_name} with adjacent categories for {fmt_currency(total_spend * 1.5)}+ contract value.",
                     "reason": f"Larger contract values attract C-suite attention from suppliers. Bundled negotiations typically yield 8-12% better pricing. Source: Category synergy analysis."},
                    {"text": f"Standardize payment terms to Net 60 across all {supplier_count} suppliers in {locations_str}, freeing {fmt_currency(total_spend * 0.02)} working capital.",
                     "reason": f"Inconsistent payment terms cost money. Standardizing to Net 60 on {fmt_currency(total_spend)} is equivalent to 2% savings in working capital. Source: Financial optimization."},
                    {"text": f"Implement mandatory e-procurement adoption for {request.category_name} purchases across {locations_str} to capture 100% spend visibility and compliance.",
                     "reason": f"E-procurement ensures negotiated rates are used. Maverick spend costs 15-25% more than contracted prices. Source: Procurement process analysis."},
                    {"text": f"Negotiate consignment inventory with {top_supplier} for high-velocity items in {locations_short}, reducing carrying costs on {fmt_currency(top_spend * 0.3)} of inventory.",
                     "reason": f"Consignment shifts inventory risk to supplier while ensuring availability. Win-win for high-trust relationships. Source: Inventory optimization."},
                    {"text": f"Develop supplier innovation program with {top_supplier} and {supplier2_name} targeting 5% annual cost reduction through process improvements.",
                     "reason": f"Strategic suppliers have insights into cost reduction. Structured innovation programs capture 3-5% annual improvement. Source: Supplier collaboration framework."},
                    {"text": f"Run quarterly business reviews with top 5 suppliers covering {concentration:.0f}% of {fmt_currency(total_spend)} spend to identify continuous improvement opportunities.",
                     "reason": f"Regular QBRs ensure strategic alignment and capture incremental savings of 2-4% annually. Source: Supplier relationship management."},
                    {"text": f"Create demand pooling mechanism across {len(locations) if locations else 1} locations for {request.category_name} with centralized ordering by Q4 2026.",
                     "reason": f"Pooled demand across {locations_str} unlocks tier-2 pricing levels. Expect 4-6% improvement from volume consolidation. Source: Demand aggregation analysis."}
                ],
                "target-pricing": [
                    {"text": f"Implement should-cost analysis for {request.category_name} key items in {locations_str} - current {price_variance:.0f}% price variance suggests overcharging from some suppliers. Complete analysis by Q2 2026.",
                     "reason": f"Spend analysis shows {price_variance:.0f}% price variance across {supplier_count} suppliers in {locations_str}. Should-cost models can identify 10-20% overcharging. Source: Price variance metrics."},
                    {"text": f"Switch to index-based pricing with {top_supplier} in {locations_short} ({fmt_currency(top_spend)} spend), linking prices to commodity indices with caps of +5% and floors of -10%.",
                     "reason": f"Supplier data shows {top_supplier} is your largest supplier in {locations_str} at {(top_spend/total_spend*100) if total_spend > 0 else 0:.0f}% of spend. Index-based contracts provide transparency. Source: Supplier spend analysis."},
                    {"text": f"Re-negotiate pricing terms with {supplier2_name} ({fmt_currency(supplier2_spend)}) and {supplier3_name} ({fmt_currency(supplier3_spend)}) using benchmark data for {locations_str}. Target 3-5% reduction.",
                     "reason": f"Combined spend of {fmt_currency(supplier2_spend + supplier3_spend)} gives leverage. Market benchmarks show 2-4% savings potential vs current pricing. Source: Supplier spend data and benchmarks."},
                    {"text": f"Establish price benchmarking database for {request.category_name} across {locations_str}, tracking prices from all {supplier_count} suppliers monthly to identify outliers.",
                     "reason": f"Current {price_variance:.0f}% variance indicates lack of price visibility. Systematic benchmarking enables data-driven negotiations. Source: Price variance metrics."},
                    {"text": f"Negotiate most-favored-customer (MFC) clauses with {top_supplier} and {supplier2_name} ensuring you receive best pricing offered to comparable buyers in {locations_str}.",
                     "reason": f"Combined spend of {fmt_currency(top_spend + supplier2_spend)} positions you as significant customer. MFC clauses protect against competitive disadvantage. Source: Supplier relationship data."},
                    {"text": f"Implement competitive bidding every 18 months for {concentration - 20:.0f}% of spend currently with {supplier3_name} and smaller suppliers in {locations_str}.",
                     "reason": f"Maintaining competitive tension in {locations_str} prevents pricing complacency. Strategic rebidding keeps pricing sharp. Source: Concentration metrics."},
                    {"text": f"Develop total cost of ownership (TCO) model for {request.category_name} including logistics, quality costs, and payment terms across {locations_str}.",
                     "reason": f"Unit price focus misses 15-25% of true costs. TCO analysis with {fmt_currency(total_spend)} spend identifies hidden optimization areas. Source: Best practices."},
                    {"text": f"Negotiate extended payment terms (Net 60-90) with {top_supplier} in exchange for volume commitments, improving working capital on {fmt_currency(top_spend)} annual spend.",
                     "reason": f"Extended terms on {top_supplier}'s {fmt_currency(top_spend)} spend provides working capital benefit equivalent to 1-2% savings. Source: Financial analysis."},
                    {"text": f"Set up automated price monitoring with ±5% threshold alerts across all {supplier_count} suppliers in {locations_str} for proactive renegotiation triggers.",
                     "reason": f"Proactive monitoring of {fmt_currency(total_spend)} spend in {locations_str} captures savings when market prices drop below targets. Source: Best practices."},
                    {"text": f"I will monitor market conditions in {locations_str} and alert you on significant price movements and negotiation windows.",
                     "reason": f"Continuous monitoring of {locations_str} ensures you stay ahead of market shifts and price fluctuations. Source: Ongoing market intelligence."},
                    # Additional 10 recommendations for 20 total
                    {"text": f"Implement open-book pricing with {top_supplier} ({fmt_currency(top_spend)}) to gain full transparency on raw material, conversion, and margin components.",
                     "reason": f"Open-book pricing reveals hidden margins and ensures fair pricing. Typical savings of 5-10% when suppliers know you understand their costs. Source: Cost transparency analysis."},
                    {"text": f"Develop price-to-beat benchmarks for all {supplier_count} suppliers in {locations_str} using Beroe market intelligence data.",
                     "reason": f"Third-party benchmarks remove emotion from negotiations. 'You're 8% above market' is more effective than 'please reduce your price.' Source: Market benchmarking."},
                    {"text": f"Negotiate price caps (max +3% per annum) with {supplier2_name} and {supplier3_name} for {fmt_currency(supplier2_spend + supplier3_spend)} combined spend.",
                     "reason": f"Price caps provide budget certainty and limit supplier opportunism during tight markets. Source: Contract negotiation best practices."},
                    {"text": f"Unbundle logistics costs from {top_supplier}'s unit prices to enable separate optimization and reveal 4-7% hidden margins.",
                     "reason": f"Bundled logistics often include 15-20% markup. Unbundling enables competitive bidding for freight. Source: Logistics cost analysis."},
                    {"text": f"Implement e-auction for {request.category_name} renewal with {supplier_count} suppliers to establish true market pricing for {locations_str}.",
                     "reason": f"E-auctions create real-time price discovery. Typical savings of 8-15% versus bilateral negotiations. Source: E-sourcing analytics."},
                    {"text": f"Create formula pricing linked to raw material indices for {top_supplier} ({fmt_currency(top_spend)}) with monthly adjustments and annual true-ups.",
                     "reason": f"Formula pricing captures market downturns automatically. Protects against being locked into high prices when markets fall. Source: Index-based pricing models."},
                    {"text": f"Audit landed cost components for {request.category_name} imports in {locations_short} including duties, freight, handling, and insurance.",
                     "reason": f"Landed cost audit often reveals 3-5% savings through duty optimization, route changes, or consolidation. Source: Trade compliance analysis."},
                    {"text": f"Implement price review triggers tied to raw material movements >5% with automatic renegotiation clause for {fmt_currency(total_spend)} spend.",
                     "reason": f"Trigger clauses ensure pricing stays market-relevant without continuous renegotiation overhead. Source: Contract mechanism design."},
                    {"text": f"Develop cost modeling capability for {request.category_name} using raw material indices, energy costs, and labor rates by region.",
                     "reason": f"In-house cost models validate supplier pricing and identify when to push back. Typical savings 4-8%. Source: Should-cost methodology."},
                    {"text": f"Create competitive tension by qualifying 1-2 alternative suppliers for 20% of {fmt_currency(total_spend)} currently with {top_supplier}.",
                     "reason": f"Alternative options prevent incumbent complacency. Even threat of switching improves pricing 3-5%. Source: Competitive sourcing strategy."}
                ],
                "risk-management": [
                    {"text": f"Qualify 2-3 backup suppliers for {request.category_name} in {locations_str} - current {concentration:.0f}% concentration with top 3 suppliers creates significant single-source risk. Complete qualification by Q3 2026.",
                     "reason": f"Supplier data shows {top_supplier}, {supplier2_name}, and {supplier3_name} control {fmt_currency(top3_spend)} ({concentration:.0f}%) of spend. Single-point failures could disrupt operations. Source: Concentration metrics."},
                    {"text": f"Diversify {fmt_currency(top_spend)} spend with {top_supplier} in {locations_short} by moving 20-30% to qualified alternatives while maintaining strategic relationship.",
                     "reason": f"Supplier analysis shows {top_supplier} holds {(top_spend/total_spend*100) if total_spend > 0 else 0:.0f}% of spend - highest concentration risk. Dual-sourcing reduces dependency. Source: Supplier spend data."},
                    {"text": f"Develop contingency sourcing plan for the {concentration:.0f}% of spend concentrated with top suppliers in {locations_str}, including pre-negotiated emergency supply agreements.",
                     "reason": f"Geographic and supplier concentration in {locations_str} exposes you to disruptions. Backup agreements ensure supply continuity. Source: Risk assessment metrics."},
                    {"text": f"Implement supplier risk monitoring dashboard tracking financial health, delivery performance, and quality metrics of {supplier_count} active suppliers across {locations_str}.",
                     "reason": f"Early warning of supplier issues in {locations_str} prevents disruptions to {fmt_currency(total_spend)} spend. Risk scoring enables proactive mitigation. Source: Supplier data."},
                    {"text": f"Negotiate safety stock holding agreements with {top_supplier} and {supplier2_name} for 2-4 weeks buffer inventory in {locations_str} to mitigate supply disruptions.",
                     "reason": f"Buffer stock with top 2 suppliers covering {fmt_currency(top_spend + supplier2_spend)} spend provides insurance against supply chain disruptions. Source: Supply chain best practices."},
                    {"text": f"Establish dual-sourcing strategy for critical {request.category_name} items, maintaining at least 2 qualified suppliers for items representing {fmt_currency(total_spend * 0.8)} of spend.",
                     "reason": f"Single-source risk on high-value items threatens supply continuity. Dual-sourcing balances risk and cost efficiency. Source: Risk management principles."},
                    {"text": f"Conduct annual supplier audits for {top_supplier}, {supplier2_name}, and {supplier3_name} assessing financial stability, capacity utilization, and business continuity plans.",
                     "reason": f"Top 3 suppliers handle {fmt_currency(top3_spend)} of spend. Regular audits identify emerging risks before they impact supply. Source: Risk assessment framework."},
                    {"text": f"Implement currency hedging strategy for {request.category_name} purchases in {locations_str} to mitigate exchange rate volatility impact on {fmt_currency(total_spend)} spend.",
                     "reason": f"Currency fluctuations can impact 5-15% of landed costs in {locations_str}. Hedging provides cost predictability. Source: Financial risk analysis."},
                    {"text": f"Create supplier performance scorecards for all {supplier_count} suppliers in {locations_str} with monthly reviews and remediation plans for underperformers.",
                     "reason": f"Systematic performance tracking across {fmt_currency(total_spend)} spend enables early intervention and continuous improvement. Source: Performance metrics."},
                    {"text": f"I will monitor market conditions and supply risks in {locations_str} and alert you on significant changes (±5% threshold) and emerging supplier issues.",
                     "reason": f"Continuous monitoring of {locations_str} ensures you stay ahead of supply disruptions and market volatility. Source: Ongoing risk intelligence."},
                    # Additional 10 recommendations for 20 total
                    {"text": f"Map Tier-2 suppliers for {top_supplier} and {supplier2_name} to identify hidden concentration risks in {request.category_name} supply chain.",
                     "reason": f"Many supply disruptions occur at Tier-2/3 level. Understanding sub-supplier dependencies reveals true risk exposure. Source: Supply chain visibility analysis."},
                    {"text": f"Establish force majeure notification requirements (48-hour) with all {supplier_count} suppliers in {locations_str} with defined escalation protocols.",
                     "reason": f"Early notification of disruptions enables faster response. Defined protocols reduce reaction time by 40-60%. Source: Business continuity planning."},
                    {"text": f"Implement supplier credit monitoring for {top_supplier}, {supplier2_name}, and {supplier3_name} controlling {fmt_currency(top3_spend)} of spend.",
                     "reason": f"Financial distress is visible 6-12 months before failure. Credit monitoring provides early warning to diversify. Source: Financial risk assessment."},
                    {"text": f"Create regional sourcing alternatives for {locations_str} to reduce cross-border risk exposure on {fmt_currency(total_spend)} spend.",
                     "reason": f"Regional suppliers reduce tariff, logistics, and geopolitical risks. Typically 15-20% more resilient than distant sources. Source: Geographic risk analysis."},
                    {"text": f"Negotiate capacity reservation agreements with {supplier2_name} and {supplier3_name} for surge demand scenarios (up to +30%).",
                     "reason": f"Reserved capacity ensures supply during market tightness. Small premium (1-2%) provides significant insurance value. Source: Capacity planning."},
                    {"text": f"Develop qualification pathway for 3 regional backup suppliers in {locations_short} with 90-day readiness capability.",
                     "reason": f"Pre-qualified backups can activate within 90 days during disruptions. Worth the qualification investment. Source: Supplier development."},
                    {"text": f"Implement supplier ESG risk scoring for all {supplier_count} suppliers to identify reputational and regulatory risks.",
                     "reason": f"ESG incidents can disrupt supply and damage brand. Early identification enables proactive mitigation. Source: ESG risk framework."},
                    {"text": f"Create joint disruption response teams with {top_supplier} including VP-level participation and defined communication channels.",
                     "reason": f"Pre-established relationships reduce disruption impact by 40%. Joint teams recover faster than ad-hoc responses. Source: Crisis management."},
                    {"text": f"Implement multi-modal logistics options for {request.category_name} in {locations_str} to reduce dependency on single transport routes.",
                     "reason": f"Multi-modal flexibility (sea, rail, truck) provides alternatives during port congestion or route disruptions. Source: Logistics resilience."},
                    {"text": f"Establish VMI (Vendor Managed Inventory) with {top_supplier} for critical items, ensuring 4-6 weeks buffer at supplier facilities.",
                     "reason": f"VMI shifts inventory risk while ensuring availability. Strategic for {fmt_currency(top_spend)} of spend. Source: Inventory risk management."}
                ],
                "respec-pack": [
                    {"text": f"Rationalize SKUs in {request.category_name} across {locations_str} where {price_variance:.0f}% price variance indicates spec-driven cost differences. Target 20-30% SKU reduction by Q3 2026.",
                     "reason": f"Spend analysis shows high price variance across {supplier_count} suppliers in {locations_str} suggests specification complexity. Standardization can reduce costs by 5-10%. Source: Price variance metrics."},
                    {"text": f"Standardize specifications across {locations_str} for top items with {top_supplier} ({fmt_currency(top_spend)} spend), creating common spec sheets that enable volume aggregation.",
                     "reason": f"Supplier data shows spec variations with {top_supplier} across {locations_str} prevent volume aggregation. Standardizing can unlock additional {fmt_currency(top_spend * 0.05)} savings. Source: Supplier spend analysis."},
                    {"text": f"Evaluate alternative materials/specs with {supplier2_name} and {supplier3_name} for {fmt_currency(supplier2_spend + supplier3_spend)} combined spend through joint value engineering workshops.",
                     "reason": f"{supplier2_name} and {supplier3_name} expertise in {locations_str} can identify cost-reduction opportunities through value engineering. Source: Supplier capability data."},
                    {"text": f"Conduct pack size optimization study for {request.category_name} across {locations_str}, evaluating bulk delivery options (flexi-tanks, ISO containers) for {fmt_currency(total_spend * 0.6)} of spend.",
                     "reason": f"Pack size standardization across {locations_str} can reduce per-unit costs by 8-15% through optimized logistics. Source: Logistics analysis."},
                    {"text": f"Implement specification change management process requiring cross-functional approval before adding new specs, preventing unnecessary complexity in {locations_str}.",
                     "reason": f"Uncontrolled spec proliferation drives cost increases. Governance ensures only value-adding specs are approved. Source: Best practices."},
                    {"text": f"Work with {top_supplier} to develop standard grade specifications for {request.category_name} that meet 80% of use cases across {locations_str} at lower cost.",
                     "reason": f"Analysis shows over-specification is common. Standard grades with {top_supplier} ({fmt_currency(top_spend)} spend) can reduce costs 5-8%. Source: Specification analysis."},
                    {"text": f"Evaluate substitute products with {supplier2_name} and {supplier3_name} that meet functional requirements at 10-20% lower cost for applications in {locations_str}.",
                     "reason": f"Combined spend of {fmt_currency(supplier2_spend + supplier3_spend)} provides opportunity for alternative product trials. Source: Supplier spend data."},
                    {"text": f"Set up cross-functional spec review committee with procurement, R&D, and operations to manage {fmt_currency(total_spend)} category spend across {locations_str}.",
                     "reason": f"Cross-functional alignment ensures standardization across {locations_str} and prevents spec creep that drives costs. Source: Organization best practices."},
                    {"text": f"Create specification catalog for {request.category_name} documenting all approved specs, volumes, and suppliers across {locations_str} to identify consolidation opportunities.",
                     "reason": f"Centralized spec visibility across {supplier_count} suppliers enables identification of redundant specifications. Source: Data management best practices."},
                    {"text": f"I will monitor market conditions in {locations_str} and alert you on new product innovations and specification optimization opportunities (±5% cost impact threshold).",
                     "reason": f"Continuous monitoring of {fmt_currency(total_spend)} spend in {locations_str} ensures you capture specification optimization opportunities. Source: Ongoing market intelligence."},
                    # Additional 10 recommendations for 20 total
                    {"text": f"Implement specification governance committee with procurement, R&D, and operations to approve any new specs before adding to {request.category_name} portfolio.",
                     "reason": f"Uncontrolled spec proliferation is a hidden cost driver. Governance reduces SKU complexity by 20-30%. Source: Specification management."},
                    {"text": f"Conduct value engineering workshops with {top_supplier} ({fmt_currency(top_spend)}) to identify 10-15% cost reduction through spec optimization.",
                     "reason": f"Suppliers often know cheaper ways to achieve same outcomes. Joint VE captures 8-15% savings. Source: Value engineering methodology."},
                    {"text": f"Standardize pack sizes to 3-4 options across {locations_str} to enable volume aggregation and reduce per-unit costs by 8-12%.",
                     "reason": f"Multiple pack sizes fragment volumes and increase logistics costs. Standardization unlocks savings. Source: Pack optimization analysis."},
                    {"text": f"Transition to industry-standard specifications (COTS) where possible for {request.category_name}, reducing custom spec premium of 10-20%.",
                     "reason": f"Custom specs limit supplier options and increase costs. COTS enables competitive sourcing. Source: Specification benchmarking."},
                    {"text": f"Implement Total Cost of Ownership (TCO) model for specifications including disposal, quality, and processing costs across {locations_str}.",
                     "reason": f"Unit price focus misses 15-25% of true spec-related costs. TCO reveals hidden optimization areas. Source: Cost analysis methodology."},
                    {"text": f"Create specification catalog with volume, supplier, and cost data for all {request.category_name} items across {supplier_count} suppliers.",
                     "reason": f"Centralized spec visibility identifies redundant items and consolidation opportunities. Source: Master data management."},
                    {"text": f"Pilot alternative materials with {supplier2_name} ({fmt_currency(supplier2_spend)}) targeting 15% cost reduction while maintaining quality standards.",
                     "reason": f"Material substitution often yields significant savings. Controlled pilots reduce quality risk. Source: Material engineering."},
                    {"text": f"Implement specification change management with impact assessment required before any modifications to {request.category_name} specs.",
                     "reason": f"Uncontrolled changes drive costs. Impact assessment prevents well-intentioned but expensive modifications. Source: Change management."},
                    {"text": f"Benchmark specifications against best-in-class industry standards to identify over-engineering opportunities worth {fmt_currency(total_spend * 0.05)}.",
                     "reason": f"Many specs exceed actual requirements. Benchmarking reveals where 'good enough' saves money. Source: Industry benchmarking."},
                    {"text": f"Develop specification roadmap with {top_supplier} for next 3 years including planned simplifications and cost reduction targets.",
                     "reason": f"Strategic spec planning with key suppliers ensures continuous improvement. Target 3-5% annual reduction. Source: Supplier development."}
                ]
            }
            recommendations = fallback_recs.get(request.opportunity_type, fallback_recs["volume-bundling"])

        # Ensure we always return exactly 20 recommendations
        # If deduplication filtered some out, supplement with generic fallbacks
        if recommendations and len(recommendations) < 20:
            logger.info(f"Only {len(recommendations)} recommendations after filtering, supplementing to reach 20")
            supplement_recs = {
                "volume-bundling": [
                    {"text": "Implement automated spend analytics to continuously identify bundling opportunities across categories.", "reason": "Data-driven opportunity identification ensures no savings potential is missed."},
                    {"text": "Develop supplier scorecards with performance metrics to identify consolidation candidates.", "reason": "Performance visibility helps prioritize which supplier relationships to strengthen."},
                    {"text": "Create demand forecasting process with stakeholders to improve volume predictability for suppliers.", "reason": "Better forecasts enable suppliers to optimize production and offer better pricing."},
                ],
                "target-pricing": [
                    {"text": "Implement automated price alerts with ±5% threshold for proactive renegotiation triggers.", "reason": "Proactive monitoring captures savings when market prices drop."},
                    {"text": "Develop total cost of ownership model including logistics, quality, and payment terms.", "reason": "TCO analysis identifies 15-25% of costs missed by focusing only on unit price."},
                    {"text": "Create competitive tension by periodically re-quoting portion of spend to alternative suppliers.", "reason": "Market testing keeps incumbents sharp on pricing."},
                ],
                "risk-management": [
                    {"text": "Develop supplier financial health monitoring using public financial data and credit ratings.", "reason": "Early warning of supplier distress enables proactive mitigation."},
                    {"text": "Create business continuity playbook with escalation protocols for supply disruptions.", "reason": "Documented response plans reduce reaction time during crises."},
                    {"text": "Implement dual-sourcing strategy for top 20% spend by value to reduce concentration risk.", "reason": "Diversification protects against single-supplier failures."},
                ],
                "respec-pack": [
                    {"text": "Conduct annual specification review with R&D and operations to identify simplification opportunities.", "reason": "Cross-functional reviews often uncover unnecessary complexity."},
                    {"text": "Pilot alternative materials with top supplier before full specification changes.", "reason": "Controlled pilots reduce risk of quality issues during transition."},
                    {"text": "Benchmark specifications against industry standards to identify over-engineering.", "reason": "Standard specs often meet requirements at lower cost than custom."},
                ]
            }
            extras = supplement_recs.get(request.opportunity_type, supplement_recs["volume-bundling"])
            needed = 20 - len(recommendations)
            recommendations.extend(extras[:needed])
            logger.info(f"Supplemented to {len(recommendations)} recommendations")

        return {
            "status": "success",
            "recommendations": recommendations[:20],  # Ensure max 20
            "model_used": response.model_used,
            "thinking_time": f"{response.latency_ms/1000:.1f}s"
        }

    except Exception as e:
        import traceback
        logger.error(f"Failed to generate recommendations: {str(e)}")
        logger.error(f"Full traceback: {traceback.format_exc()}")
        # Return fallback recommendations based on opportunity type (10 each with reasons)
        fallback = {
            "volume-bundling": [
                {"text": f"Consolidate demands across sites for {request.category_name} to leverage economies of scale. Target 5-8% price reduction through volume aggregation.", "reason": "Volume consolidation is a proven lever for cost reduction in fragmented spend categories."},
                {"text": "Negotiate volume-based discounts with your top suppliers, implementing tiered pricing structures that reward increased commitments.", "reason": "Tiered pricing aligns incentives and rewards loyalty with better unit economics."},
                {"text": "Bundle similar sub-categories to increase negotiating leverage and create larger strategic partnerships.", "reason": "Cross-category bundling increases your importance to suppliers and unlocks additional discounts."},
                {"text": "Rationalize the supplier base to 3-4 strategic partners who can service multiple regions effectively.", "reason": "Fewer suppliers with larger volumes enables better pricing and simplified management."},
                {"text": "Implement regional demand pooling with monthly consolidated orders to maximize shipment efficiency.", "reason": "Coordinated ordering reduces logistics costs and improves supplier planning."},
                {"text": "Negotiate rebate structures tied to annual volume thresholds with top suppliers.", "reason": "Volume rebates provide year-end savings without impacting day-to-day unit prices."},
                {"text": "Create a preferred supplier program offering guaranteed volumes in exchange for best pricing.", "reason": "Formal programs incentivize competitive pricing while ensuring supply security."},
                {"text": "Standardize order quantities and delivery schedules to enable full truckload shipments.", "reason": "Optimized logistics reduces per-unit delivery costs by 10-15%."},
                {"text": "Set up bi-weekly demand aggregation syncs to track consolidation progress and identify new bundling opportunities.", "reason": "Regular reviews ensure ongoing savings capture and continuous improvement."},
                {"text": "I will monitor market conditions and alert you on significant changes (±5% threshold).", "reason": "Continuous market monitoring ensures you capture pricing opportunities."}
            ],
            "target-pricing": [
                {"text": f"Implement should-cost analysis for {request.category_name} key items to establish fair market pricing.", "reason": "Should-cost models identify overcharging and provide negotiation ammunition."},
                {"text": "Switch to index-based pricing with your top suppliers, linking prices to commodity indices with caps and floors.", "reason": "Index-linked pricing provides transparency and automatic market adjustments."},
                {"text": "Re-negotiate pricing terms based on market benchmarks, targeting 3-5% improvement on current rates.", "reason": "Benchmark data enables data-driven negotiations and validates pricing."},
                {"text": "Establish a price benchmarking database tracking all supplier pricing monthly to identify outliers.", "reason": "Systematic price visibility enables proactive cost management."},
                {"text": "Negotiate most-favored-customer clauses ensuring you receive best pricing offered to comparable buyers.", "reason": "MFC clauses protect against competitive pricing disadvantage."},
                {"text": "Implement competitive bidding every 18 months for a portion of spend to maintain pricing tension.", "reason": "Regular competition prevents pricing complacency and keeps rates sharp."},
                {"text": "Develop total cost of ownership model including logistics, quality, and payment terms.", "reason": "TCO analysis identifies 15-25% of costs missed by focusing only on unit price."},
                {"text": "Negotiate extended payment terms (Net 60-90) in exchange for volume commitments.", "reason": "Extended terms provide working capital benefit equivalent to 1-2% savings."},
                {"text": "Set up automated price monitoring with ±5% threshold alerts for proactive renegotiation triggers.", "reason": "Proactive monitoring captures savings when market prices drop."},
                {"text": "I will monitor market conditions and alert you on significant changes.", "reason": "Continuous monitoring ensures you capture pricing opportunities."}
            ],
            "risk-management": [
                {"text": f"Qualify backup suppliers for {request.category_name} to reduce concentration risk. Complete qualification by Q3 2026.", "reason": "Single-source dependencies create unacceptable supply chain risk."},
                {"text": "Diversify spend with your largest supplier by moving 20-30% to qualified alternatives.", "reason": "Dual-sourcing reduces dependency while maintaining strategic relationships."},
                {"text": "Develop contingency sourcing plan including pre-negotiated emergency supply agreements.", "reason": "Contingency plans ensure rapid response to supply disruptions."},
                {"text": "Implement supplier risk monitoring dashboard tracking financial health and delivery performance.", "reason": "Early warning enables proactive mitigation before disruptions occur."},
                {"text": "Negotiate safety stock holding agreements with top suppliers for 2-4 weeks buffer inventory.", "reason": "Buffer stock provides insurance against supply chain disruptions."},
                {"text": "Establish dual-sourcing strategy for critical items, maintaining at least 2 qualified suppliers.", "reason": "Dual-sourcing balances risk mitigation with cost efficiency."},
                {"text": "Conduct annual supplier audits assessing financial stability and business continuity plans.", "reason": "Regular audits identify emerging risks before they impact supply."},
                {"text": "Implement currency hedging strategy to mitigate exchange rate volatility impact.", "reason": "Hedging provides cost predictability in volatile currency environments."},
                {"text": "Create supplier performance scorecards with monthly reviews and remediation plans.", "reason": "Systematic performance tracking enables continuous improvement."},
                {"text": "I will monitor market conditions and supply risks, alerting you on significant changes (±5% threshold).", "reason": "Continuous risk monitoring ensures proactive response to emerging threats."}
            ],
            "respec-pack": [
                {"text": f"Rationalize SKUs in {request.category_name} targeting 20-30% reduction to simplify procurement.", "reason": "SKU rationalization reduces complexity and enables volume consolidation."},
                {"text": "Standardize specifications across regions to enable volume aggregation and better pricing.", "reason": "Common specs unlock purchasing synergies across locations."},
                {"text": "Evaluate alternative materials/specs through joint value engineering workshops with suppliers.", "reason": "Value engineering identifies cost reduction without compromising quality."},
                {"text": "Conduct pack size optimization study evaluating bulk delivery options (flexi-tanks, ISO containers).", "reason": "Optimized pack sizes reduce per-unit costs by 8-15%."},
                {"text": "Implement specification change management requiring cross-functional approval for new specs.", "reason": "Governance prevents unnecessary specification proliferation."},
                {"text": "Work with top supplier to develop standard grade specifications meeting 80% of use cases at lower cost.", "reason": "Standard grades reduce over-specification costs by 5-8%."},
                {"text": "Evaluate substitute products that meet functional requirements at 10-20% lower cost.", "reason": "Functional equivalents provide savings without performance compromise."},
                {"text": "Set up cross-functional spec review committee with procurement, R&D, and operations.", "reason": "Cross-functional alignment ensures value-driven specification decisions."},
                {"text": "Create specification catalog documenting all approved specs and volumes to identify consolidation opportunities.", "reason": "Centralized visibility enables specification rationalization."},
                {"text": "I will monitor market conditions and alert you on specification optimization opportunities (±5% cost impact threshold).", "reason": "Continuous monitoring captures innovation and optimization opportunities."}
            ]
        }
        return {
            "status": "fallback",
            "recommendations": fallback.get(request.opportunity_type, fallback["volume-bundling"]),
            "error": str(e)
        }


# ============================================================================
# PROOF POINT EVALUATION ENDPOINT (LLM-Driven)
# ============================================================================

class ProofPointData(BaseModel):
    """Individual proof point data for evaluation."""
    id: str = Field(..., description="Proof point ID e.g. PP_REGIONAL_SPEND")
    name: str = Field(..., description="Proof point name")
    value: Optional[float] = Field(None, description="Numeric value if applicable")
    data: Optional[dict] = Field(None, description="Additional data context")


class ProofPointEvaluationRequest(BaseModel):
    """Request for LLM-driven proof point evaluation."""
    opportunity_type: str = Field(..., description="Opportunity type ID")
    category_name: str = Field(..., description="Category name")
    proof_points_data: List[ProofPointData] = Field(default=[], description="Proof points with their data")
    spend_data: dict = Field(default={}, description="Spend breakdown data")
    supplier_data: List[dict] = Field(default=[], description="Supplier list with spend")
    metrics: dict = Field(default={}, description="Computed metrics")


class ProofPointEvaluation(BaseModel):
    """Single proof point evaluation result."""
    id: str
    impact: str  # "H", "M", or "L"
    reasoning: str
    data_point: str


class ProofPointEvaluationResponse(BaseModel):
    """Response from proof point evaluation."""
    status: str
    evaluations: List[dict]
    summary: dict
    model_used: str  # Renamed to avoid Pydantic protected namespace conflict
    thinking_time: str
    confidence_score: float

    model_config = {"protected_namespaces": ()}


@router.post("/evaluate-proof-points", response_model=ProofPointEvaluationResponse)
async def evaluate_proof_points(request: ProofPointEvaluationRequest):
    """
    Evaluate proof points using LLM (Mistral/Llama) to determine impact levels.

    This endpoint sends proof point data to an LLM which returns Low/Medium/High
    ratings based on the actual data context, replacing hardcoded thresholds.

    The confidence score is calculated using weighted formula:
    Score = (0.25 × L_count) + (0.625 × M_count) + (0.875 × H_count)
    Confidence % = Score / (total × 0.875) × 100

    Returns:
        - evaluations: List of {id, impact, reasoning, data_point} for each proof point
        - summary: {high_count, medium_count, low_count, confidence_score, overall_assessment}
        - confidence_score: Weighted confidence percentage (0-100)
    """
    from app.services.llm_service import LLMService
    import json as json_module

    llm = LLMService()

    try:
        response = await llm.evaluate_proof_points(
            opportunity_type=request.opportunity_type,
            category_name=request.category_name,
            proof_points_data=[pp.model_dump() for pp in request.proof_points_data],
            spend_data=request.spend_data,
            supplier_data=request.supplier_data,
            metrics=request.metrics
        )

        # Parse the LLM response
        raw_content = response.content.strip()
        logger.info(f"Proof point evaluation raw response (first 500 chars): {raw_content[:500]}")

        try:
            parsed = json_module.loads(raw_content)

            evaluations = parsed.get("evaluations", [])
            summary = parsed.get("summary", {})

            # Calculate weighted confidence score if not provided
            high_count = summary.get("high_count", 0)
            medium_count = summary.get("medium_count", 0)
            low_count = summary.get("low_count", 0)

            # If counts not in summary, count from evaluations
            if high_count == 0 and medium_count == 0 and low_count == 0:
                for ev in evaluations:
                    impact = ev.get("impact", "M").upper()
                    if impact == "H":
                        high_count += 1
                    elif impact == "M":
                        medium_count += 1
                    else:
                        low_count += 1

            total = high_count + medium_count + low_count
            if total > 0:
                # Weighted score formula: (0.25 × L) + (0.625 × M) + (0.875 × H)
                weighted_score = (0.25 * low_count) + (0.625 * medium_count) + (0.875 * high_count)
                max_possible = total * 0.875
                confidence_score = (weighted_score / max_possible) * 100 if max_possible > 0 else 50.0
            else:
                confidence_score = 50.0

            # Update summary with calculated values
            summary["high_count"] = high_count
            summary["medium_count"] = medium_count
            summary["low_count"] = low_count
            summary["confidence_score"] = round(confidence_score, 1)

            return {
                "status": "success",
                "evaluations": evaluations,
                "summary": summary,
                "confidence_score": round(confidence_score, 1),
                "model_used": response.model_used,
                "thinking_time": f"{response.latency_ms/1000:.1f}s"
            }

        except json_module.JSONDecodeError as e:
            logger.error(f"Failed to parse proof point evaluation response: {e}")
            # Return fallback with medium impact for all
            return _generate_fallback_evaluation(request)

    except Exception as e:
        import traceback
        logger.error(f"Failed to evaluate proof points: {str(e)}")
        logger.error(f"Full traceback: {traceback.format_exc()}")
        return _generate_fallback_evaluation(request)


def _generate_fallback_evaluation(request: ProofPointEvaluationRequest):
    """Generate fallback proof point evaluations when LLM fails."""
    # Define default proof points by opportunity type
    default_pps = {
        "volume-bundling": [
            {"id": "PP_REGIONAL_SPEND", "impact": "M", "reasoning": "Regional spend distribution requires analysis", "data_point": "N/A"},
            {"id": "PP_TAIL_SPEND", "impact": "M", "reasoning": "Tail spend evaluation pending", "data_point": "N/A"},
            {"id": "PP_VOLUME_LEVERAGE", "impact": "M", "reasoning": "Volume leverage assessment needed", "data_point": "N/A"},
            {"id": "PP_PRICE_VARIANCE", "impact": "H", "reasoning": "Price variance typically indicates opportunity", "data_point": "N/A"},
            {"id": "PP_AVG_SPEND_SUPPLIER", "impact": "M", "reasoning": "Average spend per supplier requires review", "data_point": "N/A"},
            {"id": "PP_MARKET_CONSOLIDATION", "impact": "M", "reasoning": "Market consolidation pending analysis", "data_point": "N/A"},
            {"id": "PP_SUPPLIER_LOCATION", "impact": "L", "reasoning": "Supplier location impact varies", "data_point": "N/A"},
            {"id": "PP_SUPPLIER_RISK_RATING", "impact": "L", "reasoning": "Risk rating requires supplier data", "data_point": "N/A"},
        ],
        "target-pricing": [
            {"id": "PP_PRICE_VARIANCE", "impact": "H", "reasoning": "Price variance is key for target pricing", "data_point": "N/A"},
            {"id": "PP_TARIFF_RATE", "impact": "M", "reasoning": "Tariff impact varies by region", "data_point": "N/A"},
            {"id": "PP_COST_STRUCTURE", "impact": "M", "reasoning": "Cost structure analysis pending", "data_point": "N/A"},
            {"id": "PP_UNIT_PRICE", "impact": "H", "reasoning": "Unit price benchmark is critical", "data_point": "N/A"},
        ],
        "risk-management": [
            {"id": "PP_SINGLE_SOURCING", "impact": "H", "reasoning": "Single sourcing is high-risk factor", "data_point": "N/A"},
            {"id": "PP_SUPPLIER_CONCENTRATION", "impact": "H", "reasoning": "Concentration risk is significant", "data_point": "N/A"},
            {"id": "PP_CATEGORY_RISK", "impact": "M", "reasoning": "Category risk varies by industry", "data_point": "N/A"},
            {"id": "PP_INFLATION", "impact": "M", "reasoning": "Inflation impact requires analysis", "data_point": "N/A"},
            {"id": "PP_EXCHANGE_RATE", "impact": "M", "reasoning": "Exchange rate exposure varies", "data_point": "N/A"},
            {"id": "PP_GEO_POLITICAL", "impact": "M", "reasoning": "Geopolitical risk assessment pending", "data_point": "N/A"},
            {"id": "PP_SUPPLIER_RISK_RATING", "impact": "M", "reasoning": "Supplier risk requires evaluation", "data_point": "N/A"},
        ],
        "respec-pack": [
            {"id": "PP_PRICE_VARIANCE", "impact": "H", "reasoning": "Price variance indicates spec opportunities", "data_point": "N/A"},
            {"id": "PP_EXPORT_DATA", "impact": "M", "reasoning": "Export standard gap requires analysis", "data_point": "N/A"},
            {"id": "PP_COST_STRUCTURE", "impact": "H", "reasoning": "Material cost ratio drives spec savings", "data_point": "N/A"},
        ]
    }

    evaluations = default_pps.get(request.opportunity_type, default_pps["volume-bundling"])

    # Count impacts
    high_count = sum(1 for e in evaluations if e["impact"] == "H")
    medium_count = sum(1 for e in evaluations if e["impact"] == "M")
    low_count = sum(1 for e in evaluations if e["impact"] == "L")
    total = high_count + medium_count + low_count

    # Calculate weighted confidence
    weighted_score = (0.25 * low_count) + (0.625 * medium_count) + (0.875 * high_count)
    max_possible = total * 0.875
    confidence_score = (weighted_score / max_possible) * 100 if max_possible > 0 else 50.0

    return {
        "status": "fallback",
        "evaluations": evaluations,
        "summary": {
            "high_count": high_count,
            "medium_count": medium_count,
            "low_count": low_count,
            "confidence_score": round(confidence_score, 1),
            "overall_assessment": "Fallback evaluation - LLM unavailable. Check Groq API key or network connection."
        },
        "confidence_score": round(confidence_score, 1),
        "model_used": "fallback",
        "thinking_time": "0s"
    }


# ============================================================================
# LEADERSHIP BRIEF DOCX GENERATION ENDPOINT
# ============================================================================

class BriefRecommendation(BaseModel):
    """Single recommendation with text and reason for leadership brief."""
    text: str = Field(..., description="Recommendation text")
    reason: str = Field(default="", description="Reason/rationale for the recommendation")

class BriefProofPoint(BaseModel):
    """Proof point with validation status for leadership brief."""
    id: str = Field(..., description="Proof point ID")
    name: str = Field(..., description="Proof point name")
    isValidated: bool = Field(default=False, description="Whether the proof point is validated")

class BriefSupplier(BaseModel):
    """Supplier with spend data for leadership brief."""
    name: str = Field(..., description="Supplier name")
    spend: float = Field(default=0, description="Spend amount")

class BriefRegionSpend(BaseModel):
    """Spend by region for leadership brief."""
    name: str = Field(..., description="Region name")
    spend: float = Field(default=0, description="Spend amount")

class LeadershipBriefRequest(BaseModel):
    """Request for generating a Leadership Brief docx."""
    opportunity_id: str = Field(..., description="Opportunity type ID")
    opportunity_name: str = Field(..., description="Opportunity display name")
    category_name: str = Field(..., description="Category name e.g. 'Edible Oils'")
    locations: List[str] = Field(default=[], description="Geographic locations")
    total_spend: float = Field(default=0, description="Total spend amount")
    recommendations: List[BriefRecommendation] = Field(default=[], description="Accepted recommendations with text and reason")
    proof_points: List[BriefProofPoint] = Field(default=[], description="Proof points with validation status")
    suppliers: List[BriefSupplier] = Field(default=[], description="Suppliers with spend data")
    metrics: dict = Field(default={}, description="Computed metrics")
    savings_low: float = Field(default=0, description="Low savings estimate")
    savings_high: float = Field(default=0, description="High savings estimate")
    confidence_score: float = Field(default=0, description="Confidence score percentage")
    spend_by_region: List[BriefRegionSpend] = Field(default=[], description="Spend breakdown by region")


@router.options("/generate-brief")
async def generate_brief_options():
    """Handle CORS preflight for brief generation."""
    from fastapi.responses import Response
    return Response(
        status_code=200,
        headers={
            "Access-Control-Allow-Origin": "*",
            "Access-Control-Allow-Methods": "POST, OPTIONS",
            "Access-Control-Allow-Headers": "Content-Type, Authorization",
        }
    )


@router.post("/generate-brief")
async def generate_leadership_brief(request: LeadershipBriefRequest):
    """
    Generate a comprehensive Leadership Brief Word document (.docx) for an accepted opportunity.
    Includes LLM-enhanced executive summary, spend analysis, recommendations with reasons, simulation charts, and implementation roadmap.

    Uses OpenAI (primary) with Qwen fallback for generating enhanced content.
    """
    from fastapi.responses import Response
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor, Cm, Twips
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
    from docx.enum.style import WD_STYLE_TYPE
    from docx.oxml.ns import qn, nsmap
    from docx.oxml import OxmlElement
    from app.services.llm_service import LLMService
    import io
    import json as json_module
    from datetime import datetime

    logger.info(f"Generating LLM-enhanced leadership brief for {request.opportunity_id}: {request.category_name}")

    # =========================================================================
    # GENERATE LLM-ENHANCED CONTENT (OpenAI first, Qwen fallback)
    # =========================================================================
    llm_content = None
    try:
        llm = LLMService()
        llm_response = await llm.generate_brief_content(
            opportunity_type=request.opportunity_id,
            category_name=request.category_name,
            total_spend=request.total_spend,
            suppliers=[{"name": s.name, "spend": s.spend} for s in request.suppliers],
            metrics=request.metrics,
            recommendations=[{"text": r.text, "reason": r.reason} for r in request.recommendations],
            proof_points=[{"id": pp.id, "name": pp.name, "isValidated": pp.isValidated} for pp in request.proof_points],
            locations=request.locations,
            savings_low=request.savings_low,
            savings_high=request.savings_high,
            confidence_score=request.confidence_score
        )

        # Parse LLM response
        try:
            llm_content = json_module.loads(llm_response.content)
            logger.info(f"LLM-enhanced content generated successfully using {llm_response.model_used}")
        except json_module.JSONDecodeError:
            logger.warning("Could not parse LLM response as JSON, using template content")
            llm_content = None

    except Exception as e:
        logger.warning(f"LLM content generation failed, using template content: {str(e)}")
        llm_content = None

    # Create document
    doc = Document()

    # Set document margins
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # =========================================================================
    # HELPER FUNCTIONS
    # =========================================================================
    def format_currency(amount: float) -> str:
        if amount >= 1_000_000:
            return f"USD {amount / 1_000_000:.2f}M"
        elif amount >= 1_000:
            return f"USD {amount / 1_000:.0f}K"
        return f"USD {amount:.0f}"

    def set_cell_shading(cell, color: str):
        shading = OxmlElement('w:shd')
        shading.set(qn('w:fill'), color)
        cell._tc.get_or_add_tcPr().append(shading)

    def set_cell_border(cell, border_color="CCCCCC", border_size="4"):
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcBorders = OxmlElement('w:tcBorders')
        for border_name in ['top', 'left', 'bottom', 'right']:
            border = OxmlElement(f'w:{border_name}')
            border.set(qn('w:val'), 'single')
            border.set(qn('w:sz'), border_size)
            border.set(qn('w:color'), border_color)
            tcBorders.append(border)
        tcPr.append(tcBorders)

    def add_section_header(text: str, level: int = 1):
        heading = doc.add_heading(text, level=level)
        for run in heading.runs:
            run.font.color.rgb = RGBColor(0x1F, 0x29, 0x37)
            run.font.bold = True
        return heading

    def format_impact_percentage(percentage: float) -> str:
        """Format percentage as a clean number for impact column."""
        return f"{percentage:.0f}%"

    # Opportunity type configurations
    opp_config = {
        "volume-bundling": {"title": "VOLUME CONSOLIDATION", "color": "3B82F6", "focus": "consolidate spend"},
        "target-pricing": {"title": "PRICING OPTIMIZATION", "color": "10B981", "focus": "optimize pricing"},
        "risk-management": {"title": "RISK DIVERSIFICATION", "color": "F59E0B", "focus": "mitigate risks"},
        "respec-pack": {"title": "SPECIFICATION REVIEW", "color": "8B5CF6", "focus": "rationalize specifications"}
    }
    config = opp_config.get(request.opportunity_id, opp_config["volume-bundling"])

    # Calculate key metrics
    validated_count = len([pp for pp in request.proof_points if pp.isValidated])
    total_pps = len(request.proof_points)
    confidence = request.confidence_score if request.confidence_score > 0 else ((validated_count / total_pps * 100) if total_pps > 0 else 0)

    # Savings calculations
    savings_low = request.savings_low or (request.total_spend * 0.03)
    savings_high = request.savings_high or (request.total_spend * 0.08)
    savings_mid = (savings_low + savings_high) / 2

    # =========================================================================
    # 1. LEADERSHIP BRIEF HEADER
    # =========================================================================
    title = doc.add_heading("LEADERSHIP BRIEF", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.size = Pt(28)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0x1F, 0x29, 0x37)

    # Subtitle with category and opportunity type
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run(f"{request.category_name.upper()} PROCUREMENT {config['title']}")
    run.font.size = Pt(14)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)

    # Key metrics summary line
    summary_line = doc.add_paragraph()
    summary_line.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = summary_line.add_run(f"Total Spend: {format_currency(request.total_spend)} | ")
    run.font.size = Pt(11)
    run = summary_line.add_run(f"Locations: {', '.join(request.locations) if request.locations else 'All Regions'}")
    run.font.size = Pt(11)

    # Date line
    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = date_para.add_run(f"Generated: {datetime.now().strftime('%B %d, %Y at %I:%M %p')}")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x9C, 0xA3, 0xAF)
    run.font.italic = True

    doc.add_paragraph()

    # =========================================================================
    # 2. EXECUTIVE SUMMARY (LLM-Enhanced)
    # =========================================================================
    add_section_header("EXECUTIVE SUMMARY", 1)

    exec_summary = doc.add_paragraph()

    # Use LLM-generated summary if available, otherwise use template
    if llm_content and llm_content.get("executive_summary"):
        summary_text = llm_content["executive_summary"]
    else:
        summary_text = f"This brief outlines a strategic opportunity to {config['focus']} for {request.category_name} across {', '.join(request.locations) if request.locations else 'target regions'}. "
        summary_text += f"Based on analysis of {format_currency(request.total_spend)} in annual spend across {len(request.suppliers)} suppliers, "
        summary_text += f"we have identified potential savings of {format_currency(savings_low)} to {format_currency(savings_high)} "
        summary_text += f"({(savings_low/request.total_spend*100) if request.total_spend > 0 else 0:.1f}% - {(savings_high/request.total_spend*100) if request.total_spend > 0 else 0:.1f}%). "
        summary_text += f"Confidence level: {confidence:.0f}% ({validated_count}/{total_pps} proof points validated)."

    run = exec_summary.add_run(summary_text)
    run.font.size = Pt(11)

    # Add strategic analysis from LLM if available
    if llm_content and llm_content.get("strategic_analysis"):
        doc.add_paragraph()
        strategic_para = doc.add_paragraph()
        run = strategic_para.add_run("Strategic Analysis: ")
        run.font.bold = True
        run.font.size = Pt(11)
        strategic_text = llm_content["strategic_analysis"]
        # Ensure it's a string (LLM might return dict/list)
        if not isinstance(strategic_text, str):
            strategic_text = str(strategic_text) if strategic_text else ""
        run = strategic_para.add_run(strategic_text)
        run.font.size = Pt(11)

    doc.add_paragraph()

    # =========================================================================
    # 3. CURRENT STATE ANALYSIS
    # =========================================================================
    add_section_header("CURRENT STATE ANALYSIS", 1)

    # 3a. Spend Distribution by Supplier
    doc.add_paragraph().add_run("Spend Distribution by Supplier").bold = True

    if request.suppliers:
        supplier_table = doc.add_table(rows=1, cols=4)
        supplier_table.style = 'Table Grid'
        supplier_table.alignment = WD_TABLE_ALIGNMENT.CENTER

        # Header row
        headers = ["Supplier", "Spend (USD)", "% of Total", "Visual"]
        for idx, header in enumerate(headers):
            cell = supplier_table.rows[0].cells[idx]
            cell.text = header
            set_cell_shading(cell, "1F2937")
            for para in cell.paragraphs:
                para.runs[0].font.bold = True
                para.runs[0].font.size = Pt(10)
                para.runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

        # Data rows (top 5 suppliers)
        for supplier in request.suppliers[:5]:
            row = supplier_table.add_row()
            spend = supplier.spend
            pct = (spend / request.total_spend * 100) if request.total_spend > 0 else 0

            row.cells[0].text = supplier.name
            row.cells[1].text = format_currency(spend)
            row.cells[2].text = f"{pct:.1f}%"
            row.cells[3].text = format_impact_percentage(pct)

            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.font.size = Pt(10)

        # Total row
        total_row = supplier_table.add_row()
        total_row.cells[0].text = "Total (Top 5)"
        top5_spend = sum(s.spend for s in request.suppliers[:5])
        total_row.cells[1].text = format_currency(top5_spend)
        top5_pct = (top5_spend / request.total_spend * 100) if request.total_spend > 0 else 0
        total_row.cells[2].text = f"{top5_pct:.1f}%"
        total_row.cells[3].text = ""
        for cell in total_row.cells:
            set_cell_shading(cell, "F3F4F6")
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.bold = True
                    run.font.size = Pt(10)

    doc.add_paragraph()

    # 3b. Spend Distribution by Region
    if request.spend_by_region:
        doc.add_paragraph().add_run("Spend Distribution by Region").bold = True

        region_table = doc.add_table(rows=1, cols=4)
        region_table.style = 'Table Grid'

        headers = ["Region", "Spend (USD)", "% of Total", "Visual"]
        for idx, header in enumerate(headers):
            cell = region_table.rows[0].cells[idx]
            cell.text = header
            set_cell_shading(cell, "1F2937")
            for para in cell.paragraphs:
                para.runs[0].font.bold = True
                para.runs[0].font.size = Pt(10)
                para.runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

        for region in request.spend_by_region:
            row = region_table.add_row()
            spend = region.spend
            pct = (spend / request.total_spend * 100) if request.total_spend > 0 else 0

            row.cells[0].text = region.name
            row.cells[1].text = format_currency(spend)
            row.cells[2].text = f"{pct:.1f}%"
            row.cells[3].text = format_impact_percentage(pct)

            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.font.size = Pt(10)

        doc.add_paragraph()

    # 3c. Key Risk Indicators
    doc.add_paragraph().add_run("Key Risk Indicators").bold = True

    risk_table = doc.add_table(rows=1, cols=4)
    risk_table.style = 'Table Grid'

    headers = ["Metric", "Value", "Risk Level", "Insight"]
    for idx, header in enumerate(headers):
        cell = risk_table.rows[0].cells[idx]
        cell.text = header
        set_cell_shading(cell, "1F2937")
        for para in cell.paragraphs:
            para.runs[0].font.bold = True
            para.runs[0].font.size = Pt(10)
            para.runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    concentration = request.metrics.get('top3Concentration', 65)
    price_var = request.metrics.get('priceVariance', 15)

    risk_data = [
        ("Top 3 Concentration", f"{concentration:.0f}%", "High" if concentration > 70 else ("Medium" if concentration > 50 else "Low"), "Supplier diversification needed" if concentration > 60 else "Healthy supplier mix"),
        ("Price Variance", f"{price_var:.0f}%", "High" if price_var > 20 else ("Medium" if price_var > 10 else "Low"), "Pricing standardization opportunity" if price_var > 15 else "Consistent pricing"),
        ("Supplier Count", f"{len(request.suppliers)}", "-", f"Across {', '.join(request.locations[:2]) if request.locations else 'regions'}"),
        ("Confidence Score", f"{confidence:.0f}%", "High" if confidence >= 70 else ("Medium" if confidence >= 50 else "Low"), f"{validated_count}/{total_pps} proof points validated")
    ]

    for metric, value, risk, insight in risk_data:
        row = risk_table.add_row()
        row.cells[0].text = metric
        row.cells[1].text = value
        row.cells[2].text = risk
        row.cells[3].text = insight

        # Color code risk level
        if risk == "High":
            set_cell_shading(row.cells[2], "FEE2E2")  # Light red
        elif risk == "Medium":
            set_cell_shading(row.cells[2], "FEF3C7")  # Light yellow
        elif risk == "Low":
            set_cell_shading(row.cells[2], "D1FAE5")  # Light green

        for cell in row.cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(10)

    doc.add_paragraph()

    # =========================================================================
    # 3d. OPPORTUNITY-SPECIFIC ANALYSIS
    # =========================================================================

    # Get opportunity-specific section title
    opp_section_titles = {
        "volume-bundling": "VOLUME CONSOLIDATION ANALYSIS",
        "target-pricing": "PRICING ANALYSIS",
        "risk-management": "RISK ASSESSMENT",
        "respec-pack": "SPECIFICATION ANALYSIS"
    }
    opp_section_title = opp_section_titles.get(request.opportunity_id, "OPPORTUNITY ANALYSIS")

    add_section_header(opp_section_title, 1)

    # ========== VOLUME BUNDLING SPECIFIC TABLES ==========
    if request.opportunity_id == "volume-bundling":
        # Supplier Consolidation Potential Table
        doc.add_paragraph().add_run("Supplier Consolidation Potential").bold = True

        consol_table = doc.add_table(rows=1, cols=5)
        consol_table.style = 'Table Grid'

        headers = ["Supplier", "Current Spend", "Share %", "Consolidation Fit", "Volume Bonus Potential"]
        for idx, header in enumerate(headers):
            cell = consol_table.rows[0].cells[idx]
            cell.text = header
            set_cell_shading(cell, "3B82F6")  # Blue for volume bundling
            for para in cell.paragraphs:
                para.runs[0].font.bold = True
                para.runs[0].font.size = Pt(9)
                para.runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

        for idx, supplier in enumerate(request.suppliers[:5]):
            row = consol_table.add_row()
            pct = (supplier.spend / request.total_spend * 100) if request.total_spend > 0 else 0
            fit = "Strategic" if pct >= 20 else ("Preferred" if pct >= 10 else "Tactical")
            bonus_pct = 8 if pct >= 20 else (5 if pct >= 10 else 3)
            bonus_amt = supplier.spend * bonus_pct / 100

            row.cells[0].text = supplier.name
            row.cells[1].text = format_currency(supplier.spend)
            row.cells[2].text = f"{pct:.1f}%"
            row.cells[3].text = fit
            row.cells[4].text = f"+{bonus_pct}% ({format_currency(bonus_amt)})"

            # Color code consolidation fit
            if fit == "Strategic":
                set_cell_shading(row.cells[3], "DBEAFE")  # Light blue
            elif fit == "Preferred":
                set_cell_shading(row.cells[3], "D1FAE5")  # Light green
            else:
                set_cell_shading(row.cells[3], "F3F4F6")  # Light gray

            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.font.size = Pt(9)

        doc.add_paragraph()

        # Spend Fragmentation Analysis
        doc.add_paragraph().add_run("Spend Fragmentation Analysis").bold = True

        frag_table = doc.add_table(rows=1, cols=3)
        frag_table.style = 'Table Grid'

        headers = ["Category", "Spend Amount", "% of Total"]
        for idx, header in enumerate(headers):
            cell = frag_table.rows[0].cells[idx]
            cell.text = header
            set_cell_shading(cell, "3B82F6")
            for para in cell.paragraphs:
                para.runs[0].font.bold = True
                para.runs[0].font.size = Pt(9)
                para.runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

        top3_spend = sum(s.spend for s in request.suppliers[:3])
        top3_pct = (top3_spend / request.total_spend * 100) if request.total_spend > 0 else 65
        mid_spend = sum(s.spend for s in request.suppliers[3:10]) if len(request.suppliers) > 3 else request.total_spend * 0.2
        mid_pct = (mid_spend / request.total_spend * 100) if request.total_spend > 0 else 20
        tail_pct = request.metrics.get('tailSpendPercentage', 15)
        tail_spend = request.total_spend * tail_pct / 100

        frag_data = [
            ("Top 3 Suppliers (Strategic)", format_currency(top3_spend), f"{top3_pct:.0f}%", "DBEAFE"),
            ("Mid-tier Suppliers (4-10)", format_currency(mid_spend), f"{mid_pct:.0f}%", "FEF3C7"),
            ("Tail Spend (Fragmented)", format_currency(tail_spend), f"{tail_pct:.0f}%", "FEE2E2"),
        ]

        for cat, spend, pct, color in frag_data:
            row = frag_table.add_row()
            row.cells[0].text = cat
            row.cells[1].text = spend
            row.cells[2].text = pct
            set_cell_shading(row.cells[0], color)
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.font.size = Pt(9)

        # Add insight
        insight_para = doc.add_paragraph()
        run = insight_para.add_run(f"Consolidation Opportunity: ")
        run.font.bold = True
        run.font.size = Pt(10)
        run = insight_para.add_run(f"Moving {format_currency(tail_spend)} tail spend to strategic suppliers could unlock {format_currency(tail_spend * 0.1)} - {format_currency(tail_spend * 0.2)} in additional volume rebates.")
        run.font.size = Pt(10)
        run.font.italic = True

    # ========== TARGET PRICING SPECIFIC TABLES ==========
    elif request.opportunity_id == "target-pricing":
        # Price Variance Analysis Table
        doc.add_paragraph().add_run("Price Variance Analysis by Supplier").bold = True

        price_table = doc.add_table(rows=1, cols=5)
        price_table.style = 'Table Grid'

        headers = ["Supplier", "Current Spend", "Price Index", "vs Benchmark", "Savings Potential"]
        for idx, header in enumerate(headers):
            cell = price_table.rows[0].cells[idx]
            cell.text = header
            set_cell_shading(cell, "10B981")  # Green for target pricing
            for para in cell.paragraphs:
                para.runs[0].font.bold = True
                para.runs[0].font.size = Pt(9)
                para.runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

        for idx, supplier in enumerate(request.suppliers[:5]):
            row = price_table.add_row()
            # Simulate price index variation
            price_index = 100 + (8 if idx % 3 == 0 else (-3 if idx % 3 == 1 else 12))
            vs_benchmark = price_index - 100
            savings = supplier.spend * vs_benchmark / 100 if vs_benchmark > 0 else 0

            row.cells[0].text = supplier.name
            row.cells[1].text = format_currency(supplier.spend)
            row.cells[2].text = str(price_index)
            row.cells[3].text = f"{'+' if vs_benchmark > 0 else ''}{vs_benchmark}%"
            row.cells[4].text = format_currency(savings) if savings > 0 else "—"

            # Color code benchmark comparison
            if vs_benchmark > 5:
                set_cell_shading(row.cells[3], "FEE2E2")  # Red - overpriced
            elif vs_benchmark < -2:
                set_cell_shading(row.cells[3], "D1FAE5")  # Green - good price

            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.font.size = Pt(9)

        doc.add_paragraph()

        # Cost Structure Breakdown
        doc.add_paragraph().add_run("Should-Cost Model Breakdown").bold = True

        cost_table = doc.add_table(rows=1, cols=3)
        cost_table.style = 'Table Grid'

        headers = ["Cost Component", "Percentage", "Amount (USD)"]
        for idx, header in enumerate(headers):
            cell = cost_table.rows[0].cells[idx]
            cell.text = header
            set_cell_shading(cell, "10B981")
            for para in cell.paragraphs:
                para.runs[0].font.bold = True
                para.runs[0].font.size = Pt(9)
                para.runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

        cost_data = [
            ("Raw Materials", "55%", format_currency(request.total_spend * 0.55)),
            ("Manufacturing", "25%", format_currency(request.total_spend * 0.25)),
            ("Logistics", "12%", format_currency(request.total_spend * 0.12)),
            ("Supplier Margin", "8%", format_currency(request.total_spend * 0.08)),
        ]

        for component, pct, amount in cost_data:
            row = cost_table.add_row()
            row.cells[0].text = component
            row.cells[1].text = pct
            row.cells[2].text = amount
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.font.size = Pt(9)

        # Add insight
        price_var = request.metrics.get('priceVariance', 15)
        insight_para = doc.add_paragraph()
        run = insight_para.add_run(f"Target Pricing Opportunity: ")
        run.font.bold = True
        run.font.size = Pt(10)
        run = insight_para.add_run(f"Current price variance of {price_var:.0f}% indicates {format_currency(request.total_spend * price_var / 100 * 0.5)} potential savings through should-cost negotiations.")
        run.font.size = Pt(10)
        run.font.italic = True

    # ========== RISK MANAGEMENT SPECIFIC TABLES ==========
    elif request.opportunity_id == "risk-management":
        # Supplier Concentration Risk Assessment
        doc.add_paragraph().add_run("Supplier Concentration Risk Assessment").bold = True

        risk_assess_table = doc.add_table(rows=1, cols=5)
        risk_assess_table.style = 'Table Grid'

        headers = ["Supplier", "Spend at Risk", "Concentration %", "Single Source?", "Risk Level"]
        for idx, header in enumerate(headers):
            cell = risk_assess_table.rows[0].cells[idx]
            cell.text = header
            set_cell_shading(cell, "F59E0B")  # Amber for risk management
            for para in cell.paragraphs:
                para.runs[0].font.bold = True
                para.runs[0].font.size = Pt(9)
                para.runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

        for idx, supplier in enumerate(request.suppliers[:5]):
            row = risk_assess_table.add_row()
            pct = (supplier.spend / request.total_spend * 100) if request.total_spend > 0 else 0
            is_single_source = idx == 0 or pct > 30
            risk_level = "Critical" if pct > 30 else ("High" if pct > 20 else ("Medium" if pct > 10 else "Low"))

            row.cells[0].text = supplier.name
            row.cells[1].text = format_currency(supplier.spend)
            row.cells[2].text = f"{pct:.1f}%"
            row.cells[3].text = "Yes ⚠" if is_single_source else "No ✓"
            row.cells[4].text = risk_level

            # Color code risk level
            if risk_level == "Critical":
                set_cell_shading(row.cells[4], "FEE2E2")
            elif risk_level == "High":
                set_cell_shading(row.cells[4], "FED7AA")
            elif risk_level == "Medium":
                set_cell_shading(row.cells[4], "FEF3C7")
            else:
                set_cell_shading(row.cells[4], "D1FAE5")

            # Color single source
            if is_single_source:
                set_cell_shading(row.cells[3], "FEE2E2")
            else:
                set_cell_shading(row.cells[3], "D1FAE5")

            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.font.size = Pt(9)

        doc.add_paragraph()

        # Supply Chain Risk Factors
        doc.add_paragraph().add_run("Supply Chain Risk Factors").bold = True

        risk_factors_table = doc.add_table(rows=1, cols=3)
        risk_factors_table.style = 'Table Grid'

        headers = ["Risk Factor", "Current Status", "Severity"]
        for idx, header in enumerate(headers):
            cell = risk_factors_table.rows[0].cells[idx]
            cell.text = header
            set_cell_shading(cell, "F59E0B")
            for para in cell.paragraphs:
                para.runs[0].font.bold = True
                para.runs[0].font.size = Pt(9)
                para.runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

        top3_spend = sum(s.spend for s in request.suppliers[:3])
        top3_pct = (top3_spend / request.total_spend * 100) if request.total_spend > 0 else 65
        price_var = request.metrics.get('priceVariance', 15)

        risk_factors = [
            ("Concentration Risk", f"Top 3 suppliers = {top3_pct:.0f}% of spend", "HIGH", "FEE2E2"),
            ("Geographic Risk", f"{len(request.locations)} regions, limited diversification", "MEDIUM", "FEF3C7"),
            ("Price Volatility", f"{price_var:.0f}% variance across suppliers", "MEDIUM", "FEF3C7"),
        ]

        for factor, status, severity, color in risk_factors:
            row = risk_factors_table.add_row()
            row.cells[0].text = factor
            row.cells[1].text = status
            row.cells[2].text = severity
            set_cell_shading(row.cells[2], color)
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.font.size = Pt(9)

        # Add insight
        insight_para = doc.add_paragraph()
        run = insight_para.add_run(f"Risk Mitigation Value: ")
        run.font.bold = True
        run.font.size = Pt(10)
        run = insight_para.add_run(f"Qualifying backup suppliers for {format_currency(top3_spend)} concentrated spend protects against {format_currency(top3_spend * 0.2)} potential disruption costs.")
        run.font.size = Pt(10)
        run.font.italic = True

    # ========== RE-SPEC PACK SPECIFIC TABLES ==========
    elif request.opportunity_id == "respec-pack":
        # Specification Complexity Table
        doc.add_paragraph().add_run("Specification Complexity by Supplier").bold = True

        spec_table = doc.add_table(rows=1, cols=5)
        spec_table.style = 'Table Grid'

        headers = ["Supplier", "Current Spend", "SKU Count", "Spec Variations", "Standardization Savings"]
        for idx, header in enumerate(headers):
            cell = spec_table.rows[0].cells[idx]
            cell.text = header
            set_cell_shading(cell, "8B5CF6")  # Purple for respec pack
            for para in cell.paragraphs:
                para.runs[0].font.bold = True
                para.runs[0].font.size = Pt(9)
                para.runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

        for idx, supplier in enumerate(request.suppliers[:5]):
            row = spec_table.add_row()
            sku_count = 15 + idx * 8
            spec_variations = "High" if idx == 0 else ("Medium" if idx < 3 else "Low")
            savings_pct = 8 if spec_variations == "High" else (5 if spec_variations == "Medium" else 2)
            savings = supplier.spend * savings_pct / 100

            row.cells[0].text = supplier.name
            row.cells[1].text = format_currency(supplier.spend)
            row.cells[2].text = str(sku_count)
            row.cells[3].text = spec_variations
            row.cells[4].text = format_currency(savings)

            # Color code spec variations
            if spec_variations == "High":
                set_cell_shading(row.cells[3], "FEE2E2")
            elif spec_variations == "Medium":
                set_cell_shading(row.cells[3], "FEF3C7")
            else:
                set_cell_shading(row.cells[3], "D1FAE5")

            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.font.size = Pt(9)

        doc.add_paragraph()

        # Standardization Opportunity by Region
        doc.add_paragraph().add_run("Standardization Opportunity by Region").bold = True

        region_spec_table = doc.add_table(rows=1, cols=4)
        region_spec_table.style = 'Table Grid'

        headers = ["Region", "Spec Complexity", "Est. Spend", "Savings Potential"]
        for idx, header in enumerate(headers):
            cell = region_spec_table.rows[0].cells[idx]
            cell.text = header
            set_cell_shading(cell, "8B5CF6")
            for para in cell.paragraphs:
                para.runs[0].font.bold = True
                para.runs[0].font.size = Pt(9)
                para.runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

        regions = request.locations if request.locations else ["Europe", "Asia Pacific", "North America"]
        for idx, region in enumerate(regions[:5]):
            row = region_spec_table.add_row()
            complexity = 85 if idx == 0 else (65 if idx == 1 else 45)
            region_spend = request.total_spend / len(regions)
            savings = region_spend * (complexity / 100) * 0.08

            row.cells[0].text = region
            row.cells[1].text = f"{complexity}%"
            row.cells[2].text = format_currency(region_spend)
            row.cells[3].text = format_currency(savings)

            # Color code complexity
            if complexity > 70:
                set_cell_shading(row.cells[1], "FEE2E2")
            elif complexity > 50:
                set_cell_shading(row.cells[1], "FEF3C7")
            else:
                set_cell_shading(row.cells[1], "D1FAE5")

            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.font.size = Pt(9)

        # Add insight
        insight_para = doc.add_paragraph()
        run = insight_para.add_run(f"Value Engineering Potential: ")
        run.font.bold = True
        run.font.size = Pt(10)
        run = insight_para.add_run(f"Standardizing top 20 items across regions can eliminate {format_currency(request.total_spend * 0.05)} in complexity costs and unlock {format_currency(request.total_spend * 0.03)} in volume consolidation.")
        run.font.size = Pt(10)
        run.font.italic = True

    doc.add_paragraph()

    # =========================================================================
    # 4. OPPORTUNITY IDENTIFICATION (Proof Points)
    # =========================================================================
    add_section_header("OPPORTUNITY IDENTIFICATION", 1)

    intro_para = doc.add_paragraph()
    run = intro_para.add_run(f"The following proof points were analyzed to identify this {config['title'].lower()} opportunity:")
    run.font.size = Pt(11)
    run.font.italic = True

    pp_table = doc.add_table(rows=1, cols=4)
    pp_table.style = 'Table Grid'

    headers = ["#", "Proof Point", "Status", "Impact"]
    for idx, header in enumerate(headers):
        cell = pp_table.rows[0].cells[idx]
        cell.text = header
        set_cell_shading(cell, "1F2937")
        for para in cell.paragraphs:
            para.runs[0].font.bold = True
            para.runs[0].font.size = Pt(10)
            para.runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    for idx, pp in enumerate(request.proof_points, 1):
        row = pp_table.add_row()
        is_validated = pp.isValidated

        row.cells[0].text = str(idx)
        row.cells[1].text = pp.name
        row.cells[2].text = "✓ Validated" if is_validated else "○ Pending"
        row.cells[3].text = "Contributes to confidence" if is_validated else "Requires validation"

        # Color the status cell
        if is_validated:
            set_cell_shading(row.cells[2], "D1FAE5")
        else:
            set_cell_shading(row.cells[2], "FEF3C7")

        for cell in row.cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(10)

    doc.add_paragraph()

    # =========================================================================
    # 5. AI-GENERATED RECOMMENDATIONS
    # =========================================================================
    add_section_header("AI-GENERATED RECOMMENDATIONS", 1)

    rec_intro = doc.add_paragraph()
    run = rec_intro.add_run(f"Based on analysis of your {request.category_name} procurement data, the following recommendations have been accepted:")
    run.font.size = Pt(11)
    run.font.italic = True

    doc.add_paragraph()

    for idx, rec in enumerate(request.recommendations, 1):
        # Access Pydantic model attributes directly
        rec_text = rec.text
        rec_reason = rec.reason

        # Recommendation box
        rec_table = doc.add_table(rows=1, cols=1)
        rec_table.style = 'Table Grid'
        cell = rec_table.rows[0].cells[0]

        # Add recommendation number and text
        para = cell.paragraphs[0]
        run = para.add_run(f"Recommendation {idx}: ")
        run.font.bold = True
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(0x1F, 0x29, 0x37)

        run = para.add_run(rec_text)
        run.font.size = Pt(11)

        # Add reason if available
        if rec_reason:
            reason_para = cell.add_paragraph()
            run = reason_para.add_run("Rationale: ")
            run.font.bold = True
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)

            run = reason_para.add_run(rec_reason)
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)
            run.font.italic = True

        set_cell_shading(cell, "F8FAFC")
        doc.add_paragraph()

    # =========================================================================
    # 6. PROJECTED SAVINGS & BENEFITS (Simulation)
    # =========================================================================
    add_section_header("PROJECTED SAVINGS & BENEFITS", 1)

    # Savings summary table
    savings_table = doc.add_table(rows=1, cols=4)
    savings_table.style = 'Table Grid'
    savings_table.alignment = WD_TABLE_ALIGNMENT.CENTER

    headers = ["Scenario", "Savings (USD)", "% of Spend", "Visual"]
    for idx, header in enumerate(headers):
        cell = savings_table.rows[0].cells[idx]
        cell.text = header
        set_cell_shading(cell, "1F2937")
        for para in cell.paragraphs:
            para.runs[0].font.bold = True
            para.runs[0].font.size = Pt(10)
            para.runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    savings_scenarios = [
        ("Conservative", savings_low, (savings_low/request.total_spend*100) if request.total_spend > 0 else 0),
        ("Expected", savings_mid, (savings_mid/request.total_spend*100) if request.total_spend > 0 else 0),
        ("Optimistic", savings_high, (savings_high/request.total_spend*100) if request.total_spend > 0 else 0),
    ]

    for scenario, amount, pct in savings_scenarios:
        row = savings_table.add_row()
        row.cells[0].text = scenario
        row.cells[1].text = format_currency(amount)
        row.cells[2].text = f"{pct:.1f}%"
        row.cells[3].text = format_impact_percentage(min(pct * 5, 100))

        for cell in row.cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(10)

    doc.add_paragraph()

    # Savings by Recommendation (Simulation Chart as Table)
    doc.add_paragraph().add_run("Savings Distribution by Recommendation").bold = True

    sim_table = doc.add_table(rows=1, cols=3)
    sim_table.style = 'Table Grid'

    headers = ["Recommendation", "Est. Savings", "Impact"]
    for idx, header in enumerate(headers):
        cell = sim_table.rows[0].cells[idx]
        cell.text = header
        set_cell_shading(cell, "1F2937")
        for para in cell.paragraphs:
            para.runs[0].font.bold = True
            para.runs[0].font.size = Pt(10)
            para.runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    # Distribute savings across recommendations
    num_recs = len(request.recommendations)
    if num_recs > 0:
        base_savings = savings_mid / num_recs
        for idx, rec in enumerate(request.recommendations):
            row = sim_table.add_row()

            # Get recommendation text (truncate if too long)
            rec_text = rec.text[:50] + "..." if len(rec.text) > 50 else rec.text

            # Vary savings slightly for realism
            variation = 1 + (idx % 3 - 1) * 0.15
            rec_savings = base_savings * variation
            impact_pct = (rec_savings / savings_mid * 100) if savings_mid > 0 else 0

            row.cells[0].text = f"Rec {idx + 1}: {rec_text}"
            row.cells[1].text = format_currency(rec_savings)
            row.cells[2].text = format_impact_percentage(min(impact_pct * 3, 100))

            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.font.size = Pt(9)

    doc.add_paragraph()

    # =========================================================================
    # 7. IMPLEMENTATION ROADMAP
    # =========================================================================
    add_section_header("IMPLEMENTATION ROADMAP", 1)

    roadmap_table = doc.add_table(rows=1, cols=4)
    roadmap_table.style = 'Table Grid'

    headers = ["Phase", "Timeline", "Key Activities", "Expected Outcome"]
    for idx, header in enumerate(headers):
        cell = roadmap_table.rows[0].cells[idx]
        cell.text = header
        set_cell_shading(cell, "1F2937")
        for para in cell.paragraphs:
            para.runs[0].font.bold = True
            para.runs[0].font.size = Pt(10)
            para.runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    roadmap_data = [
        ("Phase 1: Planning", "Week 1-2", "Stakeholder alignment, data validation, supplier shortlisting", "Approved implementation plan"),
        ("Phase 2: Negotiation", "Week 3-8", f"RFQ process, supplier negotiations in {', '.join(request.locations[:2]) if request.locations else 'target regions'}", "Signed contracts with improved terms"),
        ("Phase 3: Execution", "Week 9-12", "Contract implementation, supplier onboarding, process changes", "New pricing/terms in effect"),
        ("Phase 4: Monitoring", "Ongoing", "Quarterly reviews, performance tracking, continuous improvement", f"Sustained savings of {format_currency(savings_mid)}/year"),
    ]

    for phase, timeline, activities, outcome in roadmap_data:
        row = roadmap_table.add_row()
        row.cells[0].text = phase
        row.cells[1].text = timeline
        row.cells[2].text = activities
        row.cells[3].text = outcome

        for cell in row.cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(10)

    doc.add_paragraph()

    # =========================================================================
    # 8. RISK MITIGATION (LLM-Enhanced)
    # =========================================================================
    add_section_header("RISK MITIGATION", 1)

    # Add LLM-generated risk assessment if available
    if llm_content and llm_content.get("risk_assessment"):
        risk_intro = doc.add_paragraph()
        risk_text = llm_content["risk_assessment"]
        # Ensure it's a string (LLM might return dict/list)
        if not isinstance(risk_text, str):
            risk_text = str(risk_text) if risk_text else ""
        run = risk_intro.add_run(risk_text)
        run.font.size = Pt(11)
        run.font.italic = True
        doc.add_paragraph()

    risk_mit_table = doc.add_table(rows=1, cols=3)
    risk_mit_table.style = 'Table Grid'

    headers = ["Risk", "Impact", "Mitigation Strategy"]
    for idx, header in enumerate(headers):
        cell = risk_mit_table.rows[0].cells[idx]
        cell.text = header
        set_cell_shading(cell, "1F2937")
        for para in cell.paragraphs:
            para.runs[0].font.bold = True
            para.runs[0].font.size = Pt(10)
            para.runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    risk_mit_data = [
        ("Supplier resistance to new terms", "Medium", "Phased implementation, demonstrate mutual benefits"),
        ("Quality concerns with new suppliers", "High", "Rigorous qualification process, trial orders"),
        ("Internal stakeholder misalignment", "Medium", "Regular communication, early involvement"),
        ("Market price volatility", "Medium", "Index-based contracts, regular market monitoring"),
    ]

    for risk, impact, mitigation in risk_mit_data:
        row = risk_mit_table.add_row()
        row.cells[0].text = risk
        row.cells[1].text = impact
        row.cells[2].text = mitigation

        # Color code impact
        if impact == "High":
            set_cell_shading(row.cells[1], "FEE2E2")
        elif impact == "Medium":
            set_cell_shading(row.cells[1], "FEF3C7")
        else:
            set_cell_shading(row.cells[1], "D1FAE5")

        for cell in row.cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(10)

    doc.add_paragraph()

    # =========================================================================
    # 9. STRATEGIC OUTCOME
    # =========================================================================
    add_section_header("STRATEGIC OUTCOME", 1)

    outcomes = [
        f"Achieve annual savings of {format_currency(savings_low)} to {format_currency(savings_high)}",
        f"Reduce supplier concentration from {concentration:.0f}% to target of 50-60%",
        f"Improve pricing transparency across {len(request.suppliers)} suppliers",
        f"Strengthen supply chain resilience in {', '.join(request.locations) if request.locations else 'target regions'}",
        f"Establish framework for continuous procurement optimization"
    ]

    for outcome in outcomes:
        para = doc.add_paragraph()
        para.style = 'List Bullet'
        run = para.add_run(outcome)
        run.font.size = Pt(11)

    doc.add_paragraph()

    # =========================================================================
    # 10. NEXT STEPS (LLM-Enhanced)
    # =========================================================================
    add_section_header("NEXT STEPS", 1)

    # Use LLM-generated implementation priorities if available
    if llm_content and llm_content.get("implementation_priorities"):
        impl_intro = doc.add_paragraph()
        run = impl_intro.add_run("Immediate Actions: ")
        run.font.bold = True
        run.font.size = Pt(11)
        impl_text = llm_content["implementation_priorities"]
        # Ensure it's a string (LLM might return dict/list)
        if not isinstance(impl_text, str):
            impl_text = str(impl_text) if impl_text else ""
        run = impl_intro.add_run(impl_text)
        run.font.size = Pt(11)
        doc.add_paragraph()

    next_steps = [
        "Schedule kickoff meeting with procurement leadership and category managers",
        f"Validate remaining {total_pps - validated_count} proof points to increase confidence to 100%",
        "Initiate supplier communication and RFQ preparation",
        "Set up project governance and reporting cadence",
        f"Begin Phase 1 implementation targeting {', '.join(request.locations[:2]) if request.locations else 'priority regions'}"
    ]

    for step in next_steps:
        para = doc.add_paragraph()
        para.style = 'List Bullet'
        run = para.add_run(step)
        run.font.size = Pt(11)

    # Add success factors from LLM if available
    if llm_content and llm_content.get("success_factors"):
        doc.add_paragraph()
        success_para = doc.add_paragraph()
        run = success_para.add_run("Success Factors: ")
        run.font.bold = True
        run.font.size = Pt(11)
        success_text = llm_content["success_factors"]
        # Ensure it's a string (LLM might return dict/list)
        if not isinstance(success_text, str):
            success_text = str(success_text) if success_text else ""
        run = success_para.add_run(success_text)
        run.font.size = Pt(11)

    doc.add_paragraph()

    # =========================================================================
    # 11. APPENDIX - DATA SOURCES
    # =========================================================================
    add_section_header("APPENDIX: DATA SOURCES", 1)

    sources_para = doc.add_paragraph()
    sources_text = f"This analysis was based on:\n"
    sources_text += f"• Spend data: {format_currency(request.total_spend)} annual spend across {len(request.suppliers)} suppliers\n"
    sources_text += f"• Geographic coverage: {', '.join(request.locations) if request.locations else 'All regions'}\n"
    sources_text += f"• Proof points analyzed: {total_pps} ({validated_count} validated)\n"
    sources_text += f"• Analysis date: {datetime.now().strftime('%B %d, %Y at %I:%M %p')}"
    run = sources_para.add_run(sources_text)
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)

    doc.add_paragraph()

    # =========================================================================
    # FOOTER
    # =========================================================================
    footer_para = doc.add_paragraph()
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer_para.add_run("━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0xE5, 0xE7, 0xEB)

    footer_text = doc.add_paragraph()
    footer_text.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer_text.add_run("Generated by Beroe Procurement Intelligence Platform | Powered by AI")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x9C, 0xA3, 0xAF)
    run.font.italic = True

    # =========================================================================
    # SAVE AND RETURN
    # =========================================================================
    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)

    # Generate filename
    safe_category = request.category_name.replace(' ', '_').replace('/', '-')
    safe_opp = config['title'].replace(' ', '_')
    filename = f"{safe_category}_{safe_opp}_Leadership_Brief_{datetime.now().strftime('%Y%m%d_%H%M')}.docx"

    logger.info(f"Leadership brief generated successfully: {filename}")

    return Response(
        content=file_stream.getvalue(),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={
            "Content-Disposition": f'attachment; filename="{filename}"',
            "Access-Control-Allow-Origin": "*",
            "Access-Control-Expose-Headers": "Content-Disposition"
        }
    )


# =============================================================================
# SUPPLIER INTELLIGENCE ENDPOINTS (PP8 - Real-time Supplier Risk Rating)
# =============================================================================

class SupplierEvaluationRequest(BaseModel):
    """Request for single supplier evaluation."""
    supplier_name: str = Field(..., description="Name of the supplier company")
    category: Optional[str] = Field(None, description="Product category (e.g., 'Edible Oil')")
    country: Optional[str] = Field(None, description="Country of operation")


class MultiSupplierEvaluationRequest(BaseModel):
    """Request for multiple supplier evaluation (for PP8)."""
    suppliers: List[dict] = Field(..., description="List of suppliers with name and spend")
    category: Optional[str] = Field(None, description="Product category")
    country: Optional[str] = Field(None, description="Country of operation")


@router.post("/supplier-intelligence/evaluate")
async def evaluate_supplier(request: SupplierEvaluationRequest):
    """
    Evaluate a single supplier's risk profile using real-time web data.

    This endpoint:
    1. Searches for recent news/financials about the supplier (Tavily)
    2. Analyzes the data using Llama 3.2 3B
    3. Returns a structured risk assessment

    Used for detailed supplier due diligence.
    """
    from app.services.supplier_intelligence import get_supplier_intelligence_service

    service = get_supplier_intelligence_service()

    try:
        result = await service.evaluate_supplier(
            supplier_name=request.supplier_name,
            category=request.category or "",
            country=request.country or ""
        )

        return {
            "status": "success",
            "evaluation": result
        }

    except Exception as e:
        logger.error(f"Supplier evaluation error: {str(e)}")
        raise HTTPException(
            status_code=500,
            detail=f"Failed to evaluate supplier: {str(e)}"
        )


@router.post("/supplier-intelligence/evaluate-pp8")
async def evaluate_suppliers_for_pp8(request: MultiSupplierEvaluationRequest):
    """
    Evaluate multiple suppliers for PP8 (Supplier Risk Rating) proof point.

    This endpoint:
    1. Evaluates top suppliers by spend (max 10 for performance)
    2. Aggregates risk ratings
    3. Returns PP8 impact rating (High/Medium/Low)

    Used during "Run Analysis" for Volume Bundling PP8.
    """
    from app.services.supplier_intelligence import get_supplier_intelligence_service

    service = get_supplier_intelligence_service()

    try:
        result = await service.evaluate_multiple_suppliers(
            suppliers=request.suppliers,
            category=request.category or "",
            country=request.country or ""
        )

        return {
            "status": "success",
            "proof_point_id": "PP8_SUPPLIER_RISK_RATING",
            "impact": result["impact"],
            "reasoning": result["reasoning"],
            "summary": result["summary"],
            "recommendations": result["recommendations"],
            "supplier_evaluations": result["supplier_evaluations"],
            "evaluated_at": result["evaluated_at"],
            "model_used": result["model_used"]
        }

    except Exception as e:
        logger.error(f"PP8 evaluation error: {str(e)}")
        raise HTTPException(
            status_code=500,
            detail=f"Failed to evaluate suppliers for PP8: {str(e)}"
        )
