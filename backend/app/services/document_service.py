"""
Document Service - Document Processing and Analysis

Handles:
- Document upload and parsing (PDF, DOCX, TXT)
- Text extraction
- LLM-based document analysis
- Storing analysis results
"""

from typing import Dict, List, Optional, Any, BinaryIO
from dataclasses import dataclass
from datetime import datetime
from pathlib import Path
import io
import json
import re

from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy import select

from app.services.llm_service import LLMService
from app.models.document import Document, DocumentType, ProcessingStatus
from app.config import settings
import structlog

logger = structlog.get_logger()


@dataclass
class ExtractedContent:
    """Extracted content from a document."""
    text: str
    metadata: Dict[str, Any]
    page_count: int
    word_count: int


@dataclass
class DocumentAnalysisResult:
    """Complete document analysis result."""
    document_id: str
    document_name: str
    document_type: str
    summary: str
    key_terms: List[Dict[str, Any]]
    pricing_info: Dict[str, Any]
    risks: List[Dict[str, Any]]
    opportunities: List[Dict[str, Any]]
    compliance: Dict[str, Any]
    recommendations: List[Dict[str, Any]]
    raw_analysis: Dict[str, Any]
    processed_at: str


class DocumentService:
    """
    Service for document processing and analysis.

    Supports:
    - PDF documents (contracts, agreements)
    - Word documents (playbooks, policies)
    - Text files
    - Excel/CSV (handled separately via upload service)
    """

    SUPPORTED_EXTENSIONS = {
        ".pdf": "application/pdf",
        ".docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        ".doc": "application/msword",
        ".txt": "text/plain",
        ".md": "text/markdown",
    }

    def __init__(self):
        """Initialize document service."""
        self.llm_service = LLMService()

    async def process_document(
        self,
        file_content: bytes,
        file_name: str,
        document_type: str,
        category: Optional[str] = None
    ) -> ExtractedContent:
        """
        Extract text content from a document.

        Args:
            file_content: Raw file bytes
            file_name: Name of the file
            document_type: Type of document (contract, playbook, etc.)
            category: Procurement category context

        Returns:
            ExtractedContent with extracted text and metadata
        """
        ext = Path(file_name).suffix.lower()

        if ext not in self.SUPPORTED_EXTENSIONS:
            raise ValueError(f"Unsupported file type: {ext}")

        if ext == ".pdf":
            return await self._extract_pdf(file_content, file_name)
        elif ext in [".docx", ".doc"]:
            return await self._extract_docx(file_content, file_name)
        elif ext in [".txt", ".md"]:
            return await self._extract_text(file_content, file_name)
        else:
            raise ValueError(f"No extractor for file type: {ext}")

    async def analyze_document(
        self,
        file_content: bytes,
        file_name: str,
        document_type: str,
        category: Optional[str] = None,
        extraction_focus: Optional[List[str]] = None,
        session_id: Optional[str] = None,
        db: Optional[AsyncSession] = None
    ) -> DocumentAnalysisResult:
        """
        Process and analyze a document with LLM.

        Args:
            file_content: Raw file bytes
            file_name: Name of the file
            document_type: Type of document
            category: Procurement category
            extraction_focus: Specific elements to focus on
            session_id: Session for storing results
            db: Database session

        Returns:
            DocumentAnalysisResult with analysis
        """
        # Extract text
        extracted = await self.process_document(
            file_content=file_content,
            file_name=file_name,
            document_type=document_type,
            category=category
        )

        # Analyze with LLM
        analysis_response = await self.llm_service.analyze_document(
            document_content=extracted.text,
            document_type=document_type,
            category=category,
            extraction_focus=extraction_focus
        )

        # Parse LLM response
        try:
            analysis_data = json.loads(analysis_response.content)
        except json.JSONDecodeError:
            # Handle non-JSON response
            analysis_data = {
                "summary": analysis_response.content[:500],
                "key_terms": [],
                "pricing": {},
                "risks": [],
                "opportunities": [],
                "compliance": {},
                "recommendations": []
            }

        result = DocumentAnalysisResult(
            document_id=f"doc_{datetime.utcnow().strftime('%Y%m%d%H%M%S')}",
            document_name=file_name,
            document_type=document_type,
            summary=analysis_data.get("summary", ""),
            key_terms=analysis_data.get("key_terms", []),
            pricing_info=analysis_data.get("pricing", {}),
            risks=analysis_data.get("risks", []),
            opportunities=analysis_data.get("opportunities", []),
            compliance=analysis_data.get("compliance", {}),
            recommendations=analysis_data.get("recommendations", []),
            raw_analysis=analysis_data,
            processed_at=datetime.utcnow().isoformat()
        )

        # Store in database if session provided
        if db and session_id:
            await self._store_document_analysis(db, session_id, file_name, document_type, extracted, result)

        return result

    async def analyze_multiple_documents(
        self,
        documents: List[Dict[str, Any]],
        category: Optional[str] = None,
        session_id: Optional[str] = None,
        db: Optional[AsyncSession] = None
    ) -> List[DocumentAnalysisResult]:
        """
        Analyze multiple documents and cross-reference findings.

        Args:
            documents: List of {file_content, file_name, document_type}
            category: Procurement category
            session_id: Session for storing
            db: Database session

        Returns:
            List of DocumentAnalysisResult
        """
        results = []

        for doc in documents:
            result = await self.analyze_document(
                file_content=doc["file_content"],
                file_name=doc["file_name"],
                document_type=doc.get("document_type", "other"),
                category=category,
                session_id=session_id,
                db=db
            )
            results.append(result)

        return results

    async def get_document_insights_for_category(
        self,
        session_id: str,
        category: str,
        db: AsyncSession
    ) -> Dict[str, Any]:
        """
        Get aggregated document insights for a category.

        Args:
            session_id: Session ID
            category: Category name
            db: Database session

        Returns:
            Aggregated insights from all documents
        """
        # Get documents for session
        result = await db.execute(
            select(Document)
            .where(
                Document.session_id == session_id,
                Document.category == category,
                Document.processing_status == ProcessingStatus.COMPLETED
            )
        )
        documents = result.scalars().all()

        if not documents:
            return {
                "has_documents": False,
                "message": "No analyzed documents found for this category"
            }

        # Aggregate insights
        all_risks = []
        all_opportunities = []
        all_terms = []
        all_recommendations = []

        for doc in documents:
            analysis = doc.analysis_result or {}
            all_risks.extend(analysis.get("risks", []))
            all_opportunities.extend(analysis.get("opportunities", []))
            all_terms.extend(analysis.get("key_terms", []))
            all_recommendations.extend(analysis.get("recommendations", []))

        # Sort by severity/priority
        all_risks.sort(key=lambda x: {"high": 0, "medium": 1, "low": 2}.get(x.get("severity", "low"), 2))
        all_recommendations.sort(key=lambda x: {"high": 0, "medium": 1, "low": 2}.get(x.get("priority", "low"), 2))

        return {
            "has_documents": True,
            "document_count": len(documents),
            "top_risks": all_risks[:5],
            "top_opportunities": all_opportunities[:5],
            "key_terms": all_terms[:10],
            "priority_recommendations": all_recommendations[:5],
            "documents_analyzed": [
                {
                    "name": doc.file_name,
                    "type": doc.document_type.value if doc.document_type else "unknown",
                    "analyzed_at": doc.processed_at.isoformat() if doc.processed_at else None
                }
                for doc in documents
            ]
        }

    async def _extract_pdf(self, content: bytes, file_name: str) -> ExtractedContent:
        """Extract text from PDF."""
        try:
            import pypdf
            reader = pypdf.PdfReader(io.BytesIO(content))

            text_parts = []
            for page in reader.pages:
                text_parts.append(page.extract_text() or "")

            full_text = "\n\n".join(text_parts)

            return ExtractedContent(
                text=full_text,
                metadata={
                    "file_name": file_name,
                    "format": "pdf",
                    "num_pages": len(reader.pages)
                },
                page_count=len(reader.pages),
                word_count=len(full_text.split())
            )

        except ImportError:
            logger.warning("pypdf not installed, using fallback extraction")
            return await self._fallback_extraction(content, file_name, "pdf")
        except Exception as e:
            logger.error(f"PDF extraction failed: {e}")
            return await self._fallback_extraction(content, file_name, "pdf")

    async def _extract_docx(self, content: bytes, file_name: str) -> ExtractedContent:
        """Extract text from DOCX."""
        try:
            from docx import Document as DocxDocument
            doc = DocxDocument(io.BytesIO(content))

            text_parts = []
            for para in doc.paragraphs:
                text_parts.append(para.text)

            # Also extract from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text for cell in row.cells)
                    text_parts.append(row_text)

            full_text = "\n".join(text_parts)

            return ExtractedContent(
                text=full_text,
                metadata={
                    "file_name": file_name,
                    "format": "docx",
                    "paragraphs": len(doc.paragraphs),
                    "tables": len(doc.tables)
                },
                page_count=1,  # DOCX doesn't have clear page breaks
                word_count=len(full_text.split())
            )

        except ImportError:
            logger.warning("python-docx not installed, using fallback extraction")
            return await self._fallback_extraction(content, file_name, "docx")
        except Exception as e:
            logger.error(f"DOCX extraction failed: {e}")
            return await self._fallback_extraction(content, file_name, "docx")

    async def _extract_text(self, content: bytes, file_name: str) -> ExtractedContent:
        """Extract text from plain text file."""
        # Try different encodings
        encodings = ["utf-8", "latin-1", "cp1252"]

        for encoding in encodings:
            try:
                text = content.decode(encoding)
                return ExtractedContent(
                    text=text,
                    metadata={
                        "file_name": file_name,
                        "format": "text",
                        "encoding": encoding
                    },
                    page_count=1,
                    word_count=len(text.split())
                )
            except UnicodeDecodeError:
                continue

        # Fallback with errors ignored
        text = content.decode("utf-8", errors="ignore")
        return ExtractedContent(
            text=text,
            metadata={
                "file_name": file_name,
                "format": "text",
                "encoding": "utf-8-lossy"
            },
            page_count=1,
            word_count=len(text.split())
        )

    async def _fallback_extraction(
        self,
        content: bytes,
        file_name: str,
        format_type: str
    ) -> ExtractedContent:
        """Fallback extraction when libraries are not available."""
        # Try to extract any text-like content
        try:
            text = content.decode("utf-8", errors="ignore")
            # Clean up binary artifacts
            text = re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f\x7f-\xff]', ' ', text)
            text = re.sub(r'\s+', ' ', text).strip()
        except:
            text = f"[Could not extract text from {file_name}. Please install required libraries: pypdf for PDF, python-docx for DOCX files.]"

        return ExtractedContent(
            text=text,
            metadata={
                "file_name": file_name,
                "format": format_type,
                "extraction_method": "fallback"
            },
            page_count=1,
            word_count=len(text.split())
        )

    async def _store_document_analysis(
        self,
        db: AsyncSession,
        session_id: str,
        file_name: str,
        document_type: str,
        extracted: ExtractedContent,
        result: DocumentAnalysisResult
    ):
        """Store document and analysis in database."""
        import uuid

        # Map string type to enum
        doc_type_map = {
            "contract": DocumentType.CONTRACT,
            "playbook": DocumentType.PLAYBOOK,
            "supplier_agreement": DocumentType.CONTRACT,
            "policy": DocumentType.POLICY,
            "other": DocumentType.OTHER,
        }

        document = Document(
            id=uuid.uuid4(),
            session_id=uuid.UUID(session_id) if isinstance(session_id, str) else session_id,
            file_name=file_name,
            document_type=doc_type_map.get(document_type, DocumentType.OTHER),
            file_size=len(extracted.text.encode()),
            mime_type=self.SUPPORTED_EXTENSIONS.get(Path(file_name).suffix.lower(), "application/octet-stream"),
            processing_status=ProcessingStatus.COMPLETED,
            extracted_text=extracted.text[:50000],  # Limit stored text
            analysis_result={
                "summary": result.summary,
                "key_terms": result.key_terms,
                "pricing": result.pricing_info,
                "risks": result.risks,
                "opportunities": result.opportunities,
                "compliance": result.compliance,
                "recommendations": result.recommendations
            },
            word_count=extracted.word_count,
            page_count=extracted.page_count,
            processed_at=datetime.utcnow()
        )

        db.add(document)
        await db.commit()

        logger.info(
            "Document analysis stored",
            document_id=str(document.id),
            file_name=file_name
        )


# Convenience function
async def analyze_document(
    file_content: bytes,
    file_name: str,
    document_type: str = "other",
    category: Optional[str] = None
) -> DocumentAnalysisResult:
    """Quick document analysis without database storage."""
    service = DocumentService()
    return await service.analyze_document(
        file_content=file_content,
        file_name=file_name,
        document_type=document_type,
        category=category
    )
